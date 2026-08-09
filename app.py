from flask import Flask, render_template, request, redirect, url_for, flash, send_from_directory, jsonify, session, g
from werkzeug.security import generate_password_hash, check_password_hash
import sqlite3
import os
import sys
import shutil
import csv
import io
import re
import uuid
import secrets
import hashlib
import string
import logging
from datetime import datetime
from bs4 import BeautifulSoup
import nh3
import readability
import requests

# Marker string used to namespace Univer workbook data stored inside an
# attachment's `content` column when a table has been edited in-browser rather
# than backed by an uploaded file. Kept short and unambiguous to avoid clashes.
UNIVER_DATA_PREFIX = 'univer:'

# ---------------------------------------------------------------------------
# Debug logging. Writes to a rotating log file under the data dir (sibling of
# the repo) and to stderr so failures are diagnosable on PythonAnywhere, where
# stdout/stderr from the WSGI worker is the only window into startup problems.
# Controlled by the COMPANION_DEBUG env var (any non-empty value enables it);
# defaults to ON when FLASK_DEBUG is set or running outside a production WSGI.
# ---------------------------------------------------------------------------
DEBUG_ENABLED = bool(os.environ.get('COMPANION_DEBUG') or os.environ.get('FLASK_DEBUG'))
REPO_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = os.environ.get(
    'COMPANION_DATA_DIR',
    os.path.join(os.path.dirname(REPO_DIR), 'compendium-data'),
)

logger = logging.getLogger('compendium')
if DEBUG_ENABLED:
    logger.setLevel(logging.DEBUG)
else:
    logger.setLevel(logging.INFO)
    logger.propagate = False

    def _ensure_debug_logger():
        try:
            os.makedirs(DATA_DIR, exist_ok=True)
        except OSError:
            return None
        log_path = os.path.join(DATA_DIR, 'compendium.log')
        try:
            fh = logging.FileHandler(log_path, encoding='utf-8')
        except OSError:
            return None
        fh.setLevel(logging.DEBUG)
        fh.setFormatter(logging.Formatter(
            '%(asctime)s %(levelname)s [%(name)s] %(message)s'
        ))
        logger.addHandler(fh)
        logger.info('Debug logging initialized -> %s', log_path)
        return fh

    _ensure_debug_logger()

# Always also echo to stderr so the WSGI/web-server error log captures startup.
_stderr = logging.StreamHandler(sys.stderr)
_stderr.setLevel(logging.DEBUG)
_stderr.setFormatter(logging.Formatter(
    '%(asctime)s %(levelname)s [%(name)s] %(message)s'
))
logger.addHandler(_stderr)
logger.info('Compendium starting: REPO_DIR=%s DATA_DIR=%s DEBUG=%s',
            REPO_DIR, DATA_DIR, DEBUG_ENABLED)


app = Flask(__name__)
# CHANGE THIS before deploying: a stable random secret protects session cookies.
app.secret_key = os.environ.get('COMPANION_SECRET_KEY', 'compendium-dev-secret-key-change-me')
# Runtime data (database, uploads, logs) lives OUTSIDE the repo so the code can
# be updated with `git clone`/`git pull` without touching generated data. The
# expected on-disk layout is:
#   compendium/        -> this git repo (code only)
#   compendium-data/   -> database, uploads, logs (sibling of the repo)
#   compendium-venv/   -> virtual environment (sibling of the repo)
# Paths resolve to a sibling directory named "compendium-data" next to this repo.
# NOTE: REPO_DIR and DATA_DIR are defined once, near the top of the file in the
# debug-logging block, so they are available to the logger before Flask config.

app.config['UPLOAD_FOLDER'] = os.path.join(DATA_DIR, 'uploads')
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # 100MB max upload

DB_PATH = os.path.join(DATA_DIR, 'compendium.db')
logger.info('DB_PATH=%s (exists=%s)', DB_PATH, os.path.exists(DB_PATH))


@app.before_request
def _debug_log_request():
    """Log every incoming request so a silent 404/500 is traceable."""
    if DEBUG_ENABLED:
        logger.debug('>> %s %s', request.method, request.path)


@app.after_request
def _debug_log_response(response):
    if DEBUG_ENABLED:
        logger.debug('<< %s %s -> %s', request.method, request.path,
                     response.status_code)
    return response


@app.errorhandler(Exception)
def _debug_log_exception(err):
    logger.exception('Unhandled exception on %s %s: %s',
                     request.method, request.path, err)
    raise


# Asset types are grouped into a small set of render "kinds" so that cards only
# ever need to know how to draw: richtext, image, video, table, file, link.
ASSET_TYPES = ('link', 'document', 'image', 'video', 'text', 'richtext', 'table')

TYPE_TO_KIND = {
    'link': 'link',
    'text': 'richtext',
    'richtext': 'richtext',
    'image': 'image',
    'video': 'video',
    'table': 'table',
    'document': 'file',
}

# Extensions that should be treated as a table asset when uploaded as a document.
TABLE_EXTENSIONS = ('csv', 'tsv', 'xls', 'xlsx', 'ods')

# Types that carry their payload in `content` and never have an uploaded file.
CONTENT_ONLY_TYPES = ('link', 'text', 'richtext')

# Per-kind upload allowlists. Anything not listed here is rejected outright so
# that active content (.html, .svg, .xhtml, ...) can never be stored and then
# served same-origin from /uploads/.
ALLOWED_EXTENSIONS = {
    'image': ('png', 'jpg', 'jpeg', 'gif', 'webp', 'bmp', 'avif'),
    'video': ('mp4', 'webm', 'ogg', 'ogv', 'mov', 'm4v'),
    'table': TABLE_EXTENSIONS,
    'document': (
        'pdf', 'txt', 'md', 'rtf', 'doc', 'docx', 'odt',
        'ppt', 'pptx', 'odp', 'zip', 'json', 'xml',
    ) + TABLE_EXTENSIONS,
}


def _upload_extension(att_type, original_name):
    """Validate an upload's extension for the given type.

    Returns the lowercased extension, or None if it is not allowed.
    """
    if '.' not in original_name:
        return None
    ext = original_name.rsplit('.', 1)[1].lower()
    return ext if ext in ALLOWED_EXTENSIONS.get(att_type, ()) else None


def asset_kind(att_type, filename=None):
    """Map a stored attachment type (plus filename hint) to a render kind."""
    kind = TYPE_TO_KIND.get(att_type, 'file')
    if kind == 'file' and filename and '.' in filename:
        ext = filename.rsplit('.', 1)[1].lower()
        if ext in TABLE_EXTENSIONS:
            return 'table'
    return kind


# Order the kind icons are displayed in on a statement's evidence summary, so
# the badge strip looks the same for every statement regardless of the order
# the attachments happen to have been added in.
KIND_DISPLAY_ORDER = ('link', 'richtext', 'image', 'video', 'table', 'file')


def evidence_summary(attachments):
    """Condense a statement's attachments into a total plus a per-kind tally.

    The left-hand statement list only shows how much evidence is attached and
    which kinds it is, not the evidence itself -- the cards in the right-hand
    pane do that. Returns e.g.
    {'total': 3, 'kinds': [('link', 2), ('image', 1)]}.
    """
    counts = {}
    for att in attachments or ():
        kind = asset_kind(att['type'], att['filename'])
        counts[kind] = counts.get(kind, 0) + 1
    return {
        'total': sum(counts.values()),
        'kinds': [(k, counts[k]) for k in KIND_DISPLAY_ORDER if k in counts],
    }


PREVIEW_LENGTH = 140


def preview_text(text, limit=PREVIEW_LENGTH):
    """Single source of truth for card preview truncation.

    Both the Jinja card and the JS card read the result of this, so a card
    never changes its text when it is re-rendered client-side.
    """
    value = (text or '').strip()
    if len(value) <= limit:
        return value
    return value[:limit].rstrip() + '...'


def get_user_by_id(conn, user_id):
    return conn.execute('SELECT * FROM users WHERE id = ?', (user_id,)).fetchone()


@app.before_request
def load_logged_in_user():
    """Populate g.user from the session on every request."""
    user_id = session.get('user_id')
    g.user = None
    if user_id is not None:
        conn = get_db()
        try:
            g.user = get_user_by_id(conn, user_id)
        finally:
            conn.close()


def login_required(view):
    """Redirect anonymous visitors to /login, preserving their target."""
    from functools import wraps

    @wraps(view)
    def wrapped(*args, **kwargs):
        if g.user is None:
            return redirect(url_for('login', next=request.endpoint))
        return view(*args, **kwargs)
    return wrapped


def require_topic_owner(conn, topic_id):
    """Return the topic row if the current user owns it, else None."""
    if g.user is None:
        return None
    topic = conn.execute(
        'SELECT t.* FROM topics t WHERE t.id = ?',
        (topic_id,),
    ).fetchone()
    if topic is None or topic['user_id'] != g.user['id']:
        return None
    return topic


def require_folder_owner(conn, folder_id):
    """Return the folder row if the current user owns it, else None."""
    if g.user is None:
        return None
    folder = conn.execute(
        'SELECT f.* FROM folders f WHERE f.id = ?',
        (folder_id,),
    ).fetchone()
    if folder is None or folder['user_id'] != g.user['id']:
        return None
    return folder


def require_statement_owner(conn, statement_id):
    """Return the statement row if the current user owns its topic, else None."""
    if g.user is None:
        return None
    stmt = conn.execute(
        'SELECT s.*, t.user_id FROM statements s '
        'JOIN topics t ON s.topic_id = t.id WHERE s.id = ?',
        (statement_id,),
    ).fetchone()
    if stmt is None or stmt['user_id'] != g.user['id']:
        return None
    return stmt


def visibility_clause():
    """SQL fragment + params limiting domain rows to what the viewer may see.

    Domains are shared, always-visible category labels with no owner, so every
    domain is reachable by everyone (logged in or not). Visibility of the
    topics/folders *inside* a domain is decided per row in the domain route.
    """
    return ('1=1', [])


def owner_username(conn, user_id):
    """Return the username for a domain owner, or 'unknown'."""
    if user_id is None:
        return 'unknown'
    row = conn.execute('SELECT username FROM users WHERE id = ?', (user_id,)).fetchone()
    return row['username'] if row else 'unknown'


# Random credential generation for the signup page. Kept server-side so the
# browser never derives secrets from the client clock / Math.random.
def generate_username():
    return 'user-' + secrets.token_hex(4)


def generate_password(length=14):
    alphabet = string.ascii_letters + string.digits
    return ''.join(secrets.choice(alphabet) for _ in range(length))


@app.context_processor
def inject_globals():
    """Values every template can rely on (e.g. the footer copyright year)."""
    return {
        'current_year': datetime.now().year,
        'current_user': g.user,
        'logged_in': g.user is not None,
    }


def get_global_stats(conn):
    """Corpus-wide totals shared by the landing page and the dashboard.

    Both surfaces present these as the same "how big is this corpus" figure,
    so they must be computed in one place: if the definition ever gains a
    filter, the two pages would otherwise silently disagree.
    """
    return conn.execute('''
        SELECT
            (SELECT COUNT(*) FROM domains) AS domain_count,
            (SELECT COUNT(*) FROM topics) AS topic_count,
            (SELECT COUNT(*) FROM statements) AS statement_count,
            (SELECT COUNT(*) FROM attachments) AS attachment_count
    ''').fetchone()


def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    conn.execute('PRAGMA foreign_keys = ON')
    return conn

def redirect_to_topic_referrer(topic_id, statement_id=None):
    """Redirect back to the topic page, preserving the selected statement.

    Asset/statement mutations reload the page; without this the first statement
    is reselected. We stash the active statement id in the URL so the topic
    template can re-highlight and scroll to the row the user was working in.
    """
    from urllib.parse import urlparse, parse_qsl, urlencode
    referrer = request.referrer or url_for('topic', topic_id=topic_id)
    parts = urlparse(referrer)
    params = dict(parse_qsl(parts.query))
    if statement_id:
        params['stmt'] = str(statement_id)
    else:
        params.pop('stmt', None)
    query = urlencode(params)
    target = parts._replace(query=query)
    return redirect(target.geturl())


def _ensure_column(conn, cursor, table, column, definition):
    """Add `column` to `table` only if SQLite doesn't already have it.

    SQLite's ALTER TABLE ADD COLUMN is a no-op-prone operation: it cannot be
    guarded with IF NOT EXISTS, so we inspect the schema and skip when present.
    """
    cursor.execute("PRAGMA table_info(%s)" % table)
    if any(row['name'] == column for row in cursor.fetchall()):
        return
    cursor.execute("ALTER TABLE %s ADD COLUMN %s %s" % (table, column, definition))


def _backfill_positions(conn, cursor, table, parent_column):
    """Assign sequential positions to rows that still have the default 0.

    Newly added `position` columns default every existing row to 0, which would
    leave ties. Only rows that are still 0 are renumbered by primary-key order
    inside each parent group, so we never disturb an order the user set.
    """
    cursor.execute(
        "SELECT %s, GROUP_CONCAT(id) FROM %s WHERE position = 0 GROUP BY %s"
        % (parent_column, table, parent_column)
    )
    for parent_value, id_csv in cursor.fetchall():
        ids = [i for i in id_csv.split(',') if i]
        cursor.executemany(
            "UPDATE %s SET position = ? WHERE id = ?" % table,
            [(idx, rid) for idx, rid in enumerate(ids)],
        )


def init_db():
    logger.info('init_db() called; DB_PATH=%s exists_before=%s', DB_PATH, os.path.exists(DB_PATH))
    os.makedirs(os.path.dirname(DB_PATH), exist_ok=True)
    conn = get_db()
    cursor = conn.cursor()
    
    cursor.executescript('''
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT NOT NULL UNIQUE,
            password_hash TEXT NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        );

        CREATE TABLE IF NOT EXISTS domains (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL UNIQUE,
            description TEXT
        );
        
        CREATE TABLE IF NOT EXISTS topics (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            domain_id INTEGER NOT NULL,
            name TEXT NOT NULL,
            description TEXT,
            is_public INTEGER NOT NULL DEFAULT 0,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (domain_id) REFERENCES domains (id)
        );
        
        CREATE TABLE IF NOT EXISTS statements (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            topic_id INTEGER NOT NULL,
            text TEXT NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (topic_id) REFERENCES topics (id) ON DELETE CASCADE
        );
        
        CREATE TABLE IF NOT EXISTS attachments (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            statement_id INTEGER NOT NULL,
            title TEXT NOT NULL,
            type TEXT NOT NULL CHECK (type IN ('link', 'document', 'image', 'video', 'text', 'richtext', 'table')),
            content TEXT NOT NULL,
            filename TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (statement_id) REFERENCES statements (id) ON DELETE CASCADE
        );

        CREATE TABLE IF NOT EXISTS folders (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            domain_id INTEGER NOT NULL,
            parent_id INTEGER,
            name TEXT NOT NULL,
            description TEXT,
            is_public INTEGER NOT NULL DEFAULT 0,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (domain_id) REFERENCES domains (id) ON DELETE CASCADE,
            FOREIGN KEY (parent_id) REFERENCES folders (id) ON DELETE CASCADE
        );
    ''')
    
    cursor.execute("SELECT sql FROM sqlite_master WHERE type='table' AND name='attachments'")
    existing = cursor.fetchone()
    if existing and 'richtext' not in existing[0]:
        _migrate_attachments_table(conn, cursor)

    # Ensure `source_url` column exists on attachments (idempotent).
    cursor.execute("PRAGMA table_info(attachments)")
    has_source_url = any(row['name'] == 'source_url' for row in cursor.fetchall())
    if not has_source_url:
        _migrate_attachments_source_url(conn, cursor)

    # `position` drives manual ordering of statements within a topic and of
    # assets within a statement. Both are added lazily so existing databases get
    # a sensible default (0) and are renumbered by creation order on first use.
    _ensure_column(conn, cursor, 'statements', 'position', 'INTEGER DEFAULT 0')
    _ensure_column(conn, cursor, 'attachments', 'position', 'INTEGER DEFAULT 0')
    _ensure_column(conn, cursor, 'attachments', 'tags', 'TEXT')
    # Topics and folders carry their own owner (the creator). Domains are shared,
    # owner-less category labels, so `domains.user_id` has been dropped.
    _ensure_column(conn, cursor, 'topics', 'user_id', 'INTEGER')
    _ensure_column(conn, cursor, 'folders', 'user_id', 'INTEGER')
    _ensure_column(conn, cursor, 'topics', 'is_public', 'INTEGER DEFAULT 0')
    _ensure_column(conn, cursor, 'folders', 'is_public', 'INTEGER DEFAULT 0')
    _backfill_positions(conn, cursor, 'statements', 'topic_id')
    _backfill_positions(conn, cursor, 'attachments', 'statement_id')

    # Topics gain a nullable `folder_id` linking them into a domain's folder tree.
    # SQLite cannot attach a declared FK on an ALTER ADD COLUMN, so referential
    # integrity is enforced in app code (see create_topic/update_topic) and the
    # column is nullable so a topic may sit directly under a domain if needed.
    _ensure_column(conn, cursor, 'topics', 'folder_id', 'INTEGER')

    # Remove the legacy per-domain owner column; domains are shared categories.
    _drop_domains_user_id(conn, cursor)

    _migrate_domains(conn, cursor)

    cursor.execute('SELECT COUNT(*) FROM users')
    if cursor.fetchone()[0] == 0:
        # First run with auth: create the seed owner account and assign all
        # example domains to it so the existing data is not orphaned.
        cursor.execute(
            'INSERT INTO users (username, password_hash) VALUES (?, ?)',
            ('asesh', generate_password_hash('password@8981724403')),
        )
        seed_owner_id = cursor.lastrowid
    else:
        seed_owner_id = None

    cursor.execute('SELECT COUNT(*) FROM domains')
    if cursor.fetchone()[0] == 0:
        # Domains are shared, owner-less categories. Only the seed bootstrap
        # inserts them; no user may create, edit, or delete a domain.
        cursor.executemany('INSERT INTO domains (name, description) VALUES (?, ?)', [
            ('Tech, Engineering & Systems', 'Codebases, software patterns, system architecture, AI implementations.'),
            ('Quantitative & Data Science', 'Mathematical proofs, statistical models, datasets, algorithmic logic.'),
            ('Market, Business & Corporate', 'Equity research, 10-K teardowns, macro dynamics, industry analyses.'),
            ('Empirical & Natural Science', 'Physical/biological scientific studies, experimental evidence, papers.'),
            ('Policy, Law & Governance', 'Regulatory frameworks, sociopolitical structures, legal documents.'),
            ('Culture, History & Arts', 'Historical events, literary criticism, media analysis, philosophical texts.')
        ])

    # Index foreign-key child columns so cascade deletes and the per-delete
    # filename lookups are index seeks, not full table scans.
    cursor.execute('CREATE INDEX IF NOT EXISTS idx_attachments_statement_id ON attachments (statement_id)')
    cursor.execute('CREATE INDEX IF NOT EXISTS idx_statements_topic_id ON statements (topic_id)')
    cursor.execute('CREATE INDEX IF NOT EXISTS idx_topics_domain_id ON topics (domain_id)')
    # Backs the dashboard's "most recent" lists, which order by creation time
    # and take only the newest handful of rows.
    cursor.execute('CREATE INDEX IF NOT EXISTS idx_attachments_created_at ON attachments (created_at DESC)')
    cursor.execute('CREATE INDEX IF NOT EXISTS idx_topics_created_at ON topics (created_at DESC)')
    cursor.execute('CREATE INDEX IF NOT EXISTS idx_topics_folder_id ON topics (folder_id)')
    cursor.execute('CREATE INDEX IF NOT EXISTS idx_folders_parent_id ON folders (parent_id)')
    cursor.execute('CREATE INDEX IF NOT EXISTS idx_folders_domain_id ON folders (domain_id)')

    # Reassign any owner-less content (legacy rows created before ownership
    # moved to topics/folders, or a fresh seed) to the seed account `asesh` so
    # nothing is orphaned. Statements and attachments inherit ownership through
    # their topic, so only topics/folders need a direct owner here. Guarded to
    # only touch NULL rows, making it safe to run on every startup.
    cursor.execute('SELECT id FROM users WHERE username = ?', ('asesh',))
    seed_row = cursor.fetchone()
    if seed_row is not None:
        seed_owner_id = seed_row['id']
        cursor.execute(
            'UPDATE topics SET user_id = ? WHERE user_id IS NULL', (seed_owner_id,)
        )
        cursor.execute(
            'UPDATE folders SET user_id = ? WHERE user_id IS NULL', (seed_owner_id,)
        )

    conn.commit()
    try:
        counts = {
            'users': cursor.execute('SELECT COUNT(*) FROM users').fetchone()[0],
            'domains': cursor.execute('SELECT COUNT(*) FROM domains').fetchone()[0],
            'topics': cursor.execute('SELECT COUNT(*) FROM topics').fetchone()[0],
            'statements': cursor.execute('SELECT COUNT(*) FROM statements').fetchone()[0],
            'attachments': cursor.execute('SELECT COUNT(*) FROM attachments').fetchone()[0],
        }
        logger.info('init_db() complete; row counts=%s', counts)
    except Exception as _e:
        logger.warning('init_db() finished but post-commit count failed: %s', _e)
    conn.close()


def _migrate_attachments_table(conn, cursor):
    """Rebuild `attachments` to widen its `type` CHECK constraint.

    Follows SQLite's documented table-rebuild procedure. Foreign keys are
    disabled for the duration (they cannot be toggled inside a transaction),
    the swap runs inside one explicit transaction so a crash between the DROP
    and the RENAME can never leave the database without an `attachments`
    table, and any scratch table left by a previously failed run is dropped
    first so the migration is safely re-runnable. Rows that cannot satisfy the
    foreign key are quarantined in `attachments_orphaned` rather than deleted.
    """
    # PRAGMA foreign_keys is a silent no-op inside a transaction, so finish any
    # implicit one first, then take manual control of transaction boundaries.
    conn.commit()
    prior_isolation = conn.isolation_level
    conn.isolation_level = None
    cursor.execute('PRAGMA foreign_keys = OFF')
    try:
        cursor.execute('BEGIN IMMEDIATE')
        try:
            # Left over from an earlier attempt that failed partway through.
            cursor.execute('DROP TABLE IF EXISTS attachments_new')
            # Rows orphaned while FK enforcement was inert cannot be carried
            # over, but they are user data: copy them into a quarantine table
            # rather than destroying them, and record how many were moved.
            cursor.execute('''CREATE TABLE IF NOT EXISTS attachments_orphaned AS
                SELECT * FROM attachments WHERE 0''')
            cursor.execute('''INSERT INTO attachments_orphaned
                SELECT * FROM attachments
                WHERE statement_id NOT IN (SELECT id FROM statements)''')
            quarantined = cursor.rowcount
            cursor.execute('DELETE FROM attachments WHERE statement_id NOT IN (SELECT id FROM statements)')
            if quarantined:
                app.logger.warning(
                    'attachments migration: moved %d orphaned row(s) into attachments_orphaned',
                    quarantined,
                )
            cursor.execute('''CREATE TABLE attachments_new (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                statement_id INTEGER NOT NULL,
                title TEXT NOT NULL,
                type TEXT NOT NULL CHECK (type IN ('link', 'document', 'image', 'video', 'text', 'richtext', 'table')),
                content TEXT NOT NULL,
                filename TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (statement_id) REFERENCES statements (id) ON DELETE CASCADE
            )''')
            cursor.execute('INSERT INTO attachments_new SELECT id, statement_id, title, type, content, filename, created_at FROM attachments')
            cursor.execute('DROP TABLE attachments')
            cursor.execute('ALTER TABLE attachments_new RENAME TO attachments')
            # Scoped to the rebuilt table so unrelated pre-existing violations
            # elsewhere in the schema cannot abort startup.
            violations = cursor.execute("PRAGMA foreign_key_check('attachments')").fetchall()
            if violations:
                raise sqlite3.IntegrityError(f'foreign key violations after migration: {violations}')
            cursor.execute('COMMIT')
        except Exception:
            cursor.execute('ROLLBACK')
            raise
    finally:
        cursor.execute('PRAGMA foreign_keys = ON')
        conn.isolation_level = prior_isolation


def _migrate_attachments_source_url(conn, cursor):
    """Add `source_url` column to `attachments` via table rebuild.

    Follows the same rebuild pattern as `_migrate_attachments_table` so that
    the new column is added atomically and the migration is safely re-runnable.
    """
    conn.commit()
    prior_isolation = conn.isolation_level
    conn.isolation_level = None
    cursor.execute('PRAGMA foreign_keys = OFF')
    try:
        cursor.execute('BEGIN IMMEDIATE')
        try:
            cursor.execute('DROP TABLE IF EXISTS attachments_new')
            cursor.execute('''CREATE TABLE attachments_new (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                statement_id INTEGER NOT NULL,
                title TEXT NOT NULL,
                type TEXT NOT NULL CHECK (type IN ('link', 'document', 'image', 'video', 'text', 'richtext', 'table')),
                content TEXT NOT NULL,
                filename TEXT,
                source_url TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (statement_id) REFERENCES statements (id) ON DELETE CASCADE
            )''')
            cursor.execute('''
                INSERT INTO attachments_new (id, statement_id, title, type, content, filename, created_at)
                SELECT id, statement_id, title, type, content, filename, created_at FROM attachments
            ''')
            cursor.execute('DROP TABLE attachments')
            cursor.execute('ALTER TABLE attachments_new RENAME TO attachments')
            violations = cursor.execute("PRAGMA foreign_key_check('attachments')").fetchall()
            if violations:
                raise sqlite3.IntegrityError(f'foreign key violations after migration: {violations}')
            cursor.execute('COMMIT')
        except Exception:
            cursor.execute('ROLLBACK')
            raise
    finally:
        cursor.execute('PRAGMA foreign_keys = ON')
        conn.isolation_level = prior_isolation


def _migrate_domains(conn, cursor):
    """Consolidate the domain set into the six canonical domains.

    Run on every startup so existing databases converge to the current domain
    set without re-seeding. Idempotent: each change is guarded so repeat runs
    are safe, and it is a no-op when the schema already matches. Any topics in
    legacy domains are merged into the matching canonical domain before the
    legacy row is dropped, so no user data is ever destroyed.
    """
    # The six canonical domains and the legacy domains that fold into each.
    targets = {
        'Tech, Engineering & Systems': (
            'Codebases, software patterns, system architecture, AI implementations.',
            ['Computer Science', 'Literature & Society'],
        ),
        'Quantitative & Data Science': (
            'Mathematical proofs, statistical models, datasets, algorithmic logic.',
            ['Math/Statistics/Logic', 'Math/Logic/Stat', 'Literature',
             'Math/Statistics/Logic'],
        ),
        'Market, Business & Corporate': (
            'Equity research, 10-K teardowns, macro dynamics, industry analyses.',
            ['Markets & Economics', 'Markets'],
        ),
        'Empirical & Natural Science': (
            'Physical/biological scientific studies, experimental evidence, papers.',
            ['Life Sciences', 'Natural Science'],
        ),
        'Policy, Law & Governance': (
            'Regulatory frameworks, sociopolitical structures, legal documents.',
            ['Society, Culture', 'Society'],
        ),
        'Culture, History & Arts': (
            'Historical events, literary criticism, media analysis, philosophical texts.',
            ['Literature, Media', 'History', 'Arts'],
        ),
    }

    # Ensure each canonical domain exists with its current description. Domains
    # are shared, owner-less categories; they carry no user_id at all.
    for name, (description, _legacy) in targets.items():
        cursor.execute(
            "INSERT OR IGNORE INTO domains (name, description) VALUES (?, ?)",
            (name, description),
        )

    # Map every legacy source domain to its canonical target, then merge topics
    # and drop the legacy row.
    for target_name, (_description, sources) in targets.items():
        cursor.execute("SELECT id FROM domains WHERE name = ?", (target_name,))
        target_id = cursor.fetchone()[0]
        for source in sources:
            cursor.execute("SELECT id FROM domains WHERE name = ?", (source,))
            row = cursor.fetchone()
            if not row:
                continue
            source_id = row[0]
            if source_id == target_id:
                continue
            cursor.execute(
                "UPDATE topics SET domain_id = ? WHERE domain_id = ?",
                (target_id, source_id),
            )
            cursor.execute("DELETE FROM domains WHERE id = ?", (source_id,))


def _drop_domains_user_id(conn, cursor):
    """Remove the legacy `domains.user_id` owner column via table rebuild.

    Ownership now lives on topics/folders, and domains are shared, owner-less
    category labels. SQLite ALTER cannot drop a column on the shipped SQLite
    version, so rebuild the table (create new, copy id/name/description, drop
    old, rename) guarded by a PRAGMA check so it runs exactly once.
    """
    cursor.execute("PRAGMA table_info(domains)")
    if not any(row['name'] == 'user_id' for row in cursor.fetchall()):
        return
    conn.commit()
    prior_isolation = conn.isolation_level
    conn.isolation_level = None
    cursor.execute('PRAGMA foreign_keys = OFF')
    try:
        cursor.execute('BEGIN IMMEDIATE')
        try:
            cursor.execute('DROP TABLE IF EXISTS domains_new')
            cursor.execute('''CREATE TABLE domains_new (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT NOT NULL UNIQUE,
                description TEXT
            )''')
            cursor.execute(
                'INSERT INTO domains_new (id, name, description) '
                'SELECT id, name, description FROM domains'
            )
            cursor.execute('DROP TABLE domains')
            cursor.execute('ALTER TABLE domains_new RENAME TO domains')
            violations = cursor.execute("PRAGMA foreign_key_check('domains')").fetchall()
            if violations:
                raise sqlite3.IntegrityError(f'foreign key violations after domains rebuild: {violations}')
            cursor.execute('COMMIT')
        except Exception:
            cursor.execute('ROLLBACK')
            raise
    finally:
        cursor.execute('PRAGMA foreign_keys = ON')
        conn.isolation_level = prior_isolation


@app.route('/')
def index():
    """Public marketing landing page."""
    conn = get_db()
    stats = get_global_stats(conn)
    vis_cond, vis_params = visibility_clause()
    featured = conn.execute(f'''
        SELECT d.id, d.name, d.description, COUNT(t.id) as topic_count
        FROM domains d
        LEFT JOIN topics t ON d.id = t.domain_id
        WHERE {vis_cond}
        GROUP BY d.id
        ORDER BY topic_count DESC, d.name
        LIMIT 3
    ''', vis_params).fetchall()
    conn.close()
    return render_template('landing.html', stats=stats, featured=featured)


@app.route('/about')
def about():
    return render_template('about.html')


@app.route('/alldomains')
def all_domains():
    conn = get_db()
    q = request.args.get('q', '').strip()
    vis_cond, vis_params = visibility_clause()
    if q:
        domains = conn.execute(f'''
            SELECT d.*, COUNT(t.id) as topic_count
            FROM domains d
            LEFT JOIN topics t ON d.id = t.domain_id
            WHERE ({vis_cond}) AND (LOWER(d.name) LIKE LOWER(?) OR LOWER(COALESCE(d.description,'')) LIKE LOWER(?))
            GROUP BY d.id
            ORDER BY d.name
        ''', vis_params + [f"%{q}%", f"%{q}%"]).fetchall()
    else:
        domains = conn.execute(f'''
            SELECT d.*, COUNT(t.id) as topic_count 
            FROM domains d 
            LEFT JOIN topics t ON d.id = t.domain_id 
            WHERE {vis_cond}
            GROUP BY d.id 
            ORDER BY d.name
        ''', vis_params).fetchall()
    conn.close()
    return render_template('alldomains.html', domains=domains, q=q)


@app.route('/dashboard')
@login_required
def dashboard():
    """Per-user profile / dashboard scoped to the logged-in account's content."""
    conn = get_db()
    uid = g.user['id']
    stats = conn.execute('''
        SELECT
            (SELECT COUNT(*) FROM domains) AS domain_count,
            (SELECT COUNT(*) FROM topics WHERE user_id = ?) AS topic_count,
            (SELECT COUNT(*) FROM statements s
               JOIN topics t ON s.topic_id = t.id WHERE t.user_id = ?) AS statement_count,
            (SELECT COUNT(*) FROM attachments a
               JOIN statements s ON a.statement_id = s.id
               JOIN topics t ON s.topic_id = t.id WHERE t.user_id = ?) AS attachment_count
    ''', (uid, uid, uid)).fetchone()
    recent_topics = conn.execute('''
        SELECT t.id, t.name, d.name AS domain_name,
               COUNT(DISTINCT s.id) AS statement_count
        FROM topics t
        JOIN domains d ON t.domain_id = d.id
        LEFT JOIN statements s ON s.topic_id = t.id
        WHERE t.user_id = ?
        GROUP BY t.id
        ORDER BY t.created_at DESC
        LIMIT 6
    ''', (uid,)).fetchall()
    # The LIMIT is applied before the joins so only six attachment rows are
    # ever joined and sorted; ordering after the join would force a full scan
    # of `attachments` plus a temp b-tree sort on every dashboard load.
    recent_evidence = conn.execute('''
        SELECT a.title, a.type, a.filename,
               t.id AS topic_id, t.name AS topic_name
        FROM (
            SELECT a.id, a.statement_id, a.title, a.type, a.filename, a.created_at
            FROM attachments a
            JOIN statements s ON a.statement_id = s.id
            JOIN topics t ON s.topic_id = t.id
            WHERE t.user_id = ?
            ORDER BY a.created_at DESC
            LIMIT 6
        ) a
        JOIN statements s ON a.statement_id = s.id
        JOIN topics t ON s.topic_id = t.id
        ORDER BY a.created_at DESC
    ''', (uid,)).fetchall()
    # Statements and attachments are pre-aggregated per topic so the deepest
    # table is not fanned out across the whole join; counting DISTINCT over
    # the full cross-product made this scale with total attachments rather
    # than with the handful of domains actually displayed.
    domain_breakdown = conn.execute('''
        SELECT d.id, d.name,
               COUNT(DISTINCT t.id) AS topic_count,
               COALESCE(SUM(ts.statement_count), 0) AS statement_count,
               COALESCE(SUM(ts.attachment_count), 0) AS attachment_count
        FROM domains d
        LEFT JOIN topics t ON t.domain_id = d.id AND t.user_id = ?
        LEFT JOIN (
            SELECT s.topic_id,
                   COUNT(DISTINCT s.id) AS statement_count,
                   COUNT(a.id) AS attachment_count
            FROM statements s
            LEFT JOIN attachments a ON a.statement_id = s.id
            GROUP BY s.topic_id
        ) ts ON ts.topic_id = t.id
        GROUP BY d.id
        ORDER BY topic_count DESC, d.name
    ''', (uid,)).fetchall()
    conn.close()
    return render_template(
        'dashboard.html',
        stats=stats,
        recent_topics=recent_topics,
        recent_evidence=recent_evidence,
        domain_breakdown=domain_breakdown,
        asset_kind=asset_kind,
    )

@app.route('/domain/<int:domain_id>')
def domain(domain_id):
    conn = get_db()
    domain = conn.execute('SELECT * FROM domains WHERE id = ?', (domain_id,)).fetchone()
    if not domain:
        conn.close()
        flash('Domain not found')
        return redirect(url_for('all_domains'))
    # Domains are shared, always-visible category labels with no owner, so every
    # domain is reachable by everyone (logged in or not). Visibility of the
    # topics/folders inside is decided per row below.
    uid = g.user['id'] if g.user is not None else None
    if uid is not None:
        topic_vis = 'AND (t.user_id = ? OR t.is_public = 1)'
        topic_params = [uid]
        folder_vis = 'AND (f.user_id = ? OR f.is_public = 1)'
        folder_params = [uid]
    else:
        topic_vis = 'AND t.is_public = 1'
        topic_params = []
        folder_vis = 'AND f.is_public = 1'
        folder_params = []
    # Topics with no folder (shouldn't happen post-backfill, but keep the page
    # honest if a topic exists with folder_id NULL). The viewer sees their own
    # topics plus any public topic in the domain.
    loose_topics = conn.execute(f'''
        SELECT t.*,
               COUNT(DISTINCT s.id) as statement_count,
               COUNT(DISTINCT a.id) as attachment_count
        FROM topics t
        LEFT JOIN statements s ON t.id = s.topic_id
        LEFT JOIN attachments a ON s.id = a.statement_id
        WHERE t.domain_id = ? AND t.folder_id IS NULL {topic_vis}
        GROUP BY t.id
        ORDER BY t.created_at DESC
    ''', (domain_id,) + tuple(topic_params)).fetchall()
    # Folders visible to the viewer: their own folders plus any public folder,
    # plus the ancestor chain so a public topic's path renders. Logged-out users
    # only ever see public folders.
    if uid is not None:
        folder_ids = [r['id'] for r in conn.execute('''
            WITH RECURSIVE subtree(id) AS (
                SELECT f.id FROM folders f
                WHERE f.domain_id = ?
                  AND (f.user_id = ? OR f.is_public = 1)
                UNION ALL
                SELECT f.parent_id FROM folders f JOIN subtree s ON f.id = s.id
                WHERE f.parent_id IS NOT NULL
            )
            SELECT DISTINCT id FROM subtree
        ''', (domain_id, uid)).fetchall()]
    else:
        folder_ids = [r['id'] for r in conn.execute('''
            WITH RECURSIVE subtree(id) AS (
                SELECT f.id FROM folders f
                WHERE f.domain_id = ? AND f.is_public = 1
                UNION ALL
                SELECT f.parent_id FROM folders f JOIN subtree s ON f.id = s.id
                WHERE f.parent_id IS NOT NULL
            )
            SELECT DISTINCT id FROM subtree
        ''', (domain_id,)).fetchall()]
    if folder_ids:
        placeholders = ','.join('?' * len(folder_ids))
        folders = conn.execute(
            f'SELECT * FROM folders WHERE id IN ({placeholders}) ORDER BY name',
            folder_ids,
        ).fetchall()
    else:
        folders = []
    q = request.args.get('q', '').strip()
    if q:
        # Gather every topic id in this domain (root-folder subtrees + loose).
        all_topic_ids = []
        for f in folders:
            _, sub_topic_ids = get_folder_subtree(conn, f['id'])
            all_topic_ids.extend(sub_topic_ids)
        all_topic_ids.extend(t['id'] for t in loose_topics)
        matched = topic_ids_matching(conn, all_topic_ids, q)

        # A folder is kept if its own name/desc matches OR it directly contains a
        # matched topic OR any descendant folder is kept. A recursive walk over
        # the folder list preserves ancestors of a deep match (no orphans).
        children_of = {}
        folder_by_id = {}
        for f in folders:
            children_of.setdefault(f['parent_id'], []).append(f['id'])
            folder_by_id[f['id']] = f

        def folder_matches(folder_id):
            f = folder_by_id[folder_id]
            if q.lower() in (f['name'] or '').lower() \
               or q.lower() in (f['description'] or '').lower():
                return True
            for child_id in children_of.get(folder_id, []):
                if folder_matches(child_id):
                    return True
            return False

        kept_ids = set()
        for f in folders:
            fid = f['id']
            name_desc_match = (q.lower() in (f['name'] or '').lower()
                               or q.lower() in (f['description'] or '').lower())
            contains_matched = any(tid in matched for tid in
                                  get_folder_subtree(conn, fid)[1])
            if name_desc_match or contains_matched or folder_matches(fid):
                kept_ids.add(fid)
        folders = [f for f in folders if f['id'] in kept_ids]
        loose_topics = [t for t in loose_topics if t['id'] in matched
                        or q.lower() in (t['name'] or '').lower()
                        or q.lower() in (t['description'] or '').lower()]
    folder_tree = build_folder_tree(conn, folders, topic_vis, topic_params)
    # Flat, depth-indented folder list for the Move-topic <select>. Walk the
    # nested tree so parent folders always precede their children.
    all_folders = []
    def walk(nodes, depth):
        for n in nodes:
            all_folders.append({'id': n['id'], 'name': n['name'], 'depth': depth, 'indent': '\u00a0' * (depth * 4)})
            walk(n['children'], depth + 1)
    walk(folder_tree, 0)
    conn.close()
    return render_template(
        'domain.html',
        domain=domain,
        folder_tree=folder_tree,
        all_folders=all_folders,
        loose_topics=loose_topics,
        q=q,
    )

@app.route('/topic/<int:topic_id>')
def topic(topic_id):
    conn = get_db()
    topic = conn.execute('''
        SELECT t.*, d.name as domain_name
        FROM topics t 
        JOIN domains d ON t.domain_id = d.id 
        WHERE t.id = ?
    ''', (topic_id,)).fetchone()
    # Access: owner sees it; everyone else only when the topic is public.
    if topic and topic['user_id'] != (g.user['id'] if g.user else None):
        if not topic['is_public']:
            conn.close()
            flash('This topic is private')
            return redirect(url_for('all_domains'))
    can_edit = topic is not None and topic['user_id'] == (g.user['id'] if g.user else None)
    owner_name = owner_username(conn, topic['user_id']) if topic else 'unknown'
    # Walk the folder's ancestor chain (parent_id upward) to build the breadcrumb
    # path domain › … › folder. Computed here so the template only renders it.
    folder_path = []
    if topic and topic['folder_id']:
        cursor = conn.cursor()
        fid = topic['folder_id']
        seen = set()
        while fid and fid not in seen:
            seen.add(fid)
            row = cursor.execute(
                'SELECT id, parent_id, name FROM folders WHERE id = ?', (fid,)
            ).fetchone()
            if not row:
                break
            folder_path.insert(0, {'id': row['id'], 'name': row['name']})
            fid = row['parent_id']
    # Flat, depth-indented folder list for the Move-topic <select> (the topic
    # page header also exposes Move). Reuses the same nesting as the domain page.
    all_folders = []
    if topic:
        domain_folders = conn.execute(
            'SELECT * FROM folders WHERE domain_id = ? ORDER BY name', (topic['domain_id'],)
        ).fetchall()
        domain_tree = build_folder_tree(conn, domain_folders)

        def walk(nodes, depth):
            for n in nodes:
                all_folders.append({'id': n['id'], 'name': n['name'], 'depth': depth, 'indent': '\u00a0' * (depth * 4)})
                walk(n['children'], depth + 1)
        walk(domain_tree, 0)
    # All topics in this domain, used by the statement "Move to topic" modal. The
    # current topic is excluded client-side (a statement cannot move to itself).
    domain_topics = []
    if topic:
        domain_topics = conn.execute(
            'SELECT id, name FROM topics WHERE domain_id = ? ORDER BY name',
            (topic['domain_id'],),
        ).fetchall()
    statements = conn.execute('SELECT * FROM statements WHERE topic_id = ? ORDER BY position ASC, created_at ASC', (topic_id,)).fetchall()
    statement_ids = [s['id'] for s in statements]
    if statement_ids:
        placeholders = ','.join('?' * len(statement_ids))
        attachments = conn.execute(f'SELECT * FROM attachments WHERE statement_id IN ({placeholders}) ORDER BY position ASC, created_at DESC', statement_ids).fetchall()
    else:
        attachments = []
    q = request.args.get('q', '').strip()
    if q:
        pat = f"%{q}%"
        # A statement matches if its own text matches, or the topic name/desc
        # matches (show all its attachments either way).
        topic_match = (q.lower() in (topic['name'] or '').lower()
                       or q.lower() in (topic['description'] or '').lower())
        stmt_match = {s['id'] for s in statements if pat.lower() in (s['text'] or '').lower()}
        keep_stmt_ids = stmt_match | ({s['id'] for s in statements} if topic_match else set())
        # Attachments: keep all for kept statements unless only an attachment title
        # matched (then keep just the matching attachment).
        att_match = {a['id'] for a in attachments if pat.lower() in (a['title'] or '').lower()}
        statements = [s for s in statements if s['id'] in keep_stmt_ids]
        statement_ids = [s['id'] for s in statements]
        if statement_ids:
            placeholders = ','.join('?' * len(statement_ids))
            filtered_ids = set(statement_ids)
            attachments = [a for a in attachments
                           if a['statement_id'] in filtered_ids
                           and (a['statement_id'] in stmt_match or a['id'] in att_match
                                or topic_match)]
        else:
            attachments = []
    attachments_by_statement = {}
    for att in attachments:
        attachments_by_statement.setdefault(att['statement_id'], []).append(att)

    # Which statement the evidence pane should show on load. Driven by ?stmt=
    # so a reload after editing/adding an asset stays on the row the user was
    # working in, instead of snapping back to the first statement.
    raw_stmt = request.args.get('stmt')
    active_statement_id = None
    if raw_stmt and raw_stmt.isdigit():
        sid = int(raw_stmt)
        if any(s['id'] == sid for s in statements):
            active_statement_id = sid
    if active_statement_id is None and statements:
        active_statement_id = statements[0]['id']
    # Plain-dict version, keyed by string id, for JSON injection into the page so
    # the right-hand evidence pane can re-render on statement click. For text
    # kinds (link/richtext) the full body is large, so only the preview ships;
    # the modal/edit form fetch the full body on demand via /attachment/<id>.
    # For file-backed kinds the "content" is just a short URL, so it is kept.
    evidence_data = {
        str(sid): [
            {
                'id': a['id'],
                'statement_id': a['statement_id'],
                'title': a['title'],
                'type': a['type'],
                'kind': asset_kind(a['type'], a['filename']),
                'preview': preview_text(a['content']),
                'filename': a['filename'] or '',
                'tags': a['tags'] or '',
                **({'content': a['content']} if a['type'] not in CONTENT_ONLY_TYPES else {}),
            }
            for a in atts
        ]
        for sid, atts in attachments_by_statement.items()
    }
    conn.close()
    if not topic:
        flash('Topic not found')
        return redirect(url_for('all_domains'))
    return render_template(
        'topic.html',
        topic=topic,
        folder_path=folder_path,
        all_folders=all_folders,
        domain_topics=domain_topics,
        statements=statements,
        attachments_by_statement=attachments_by_statement,
        evidence_data=evidence_data,
        evidence_summary=evidence_summary,
        asset_kind=asset_kind,
        preview_text=preview_text,
        CONTENT_ONLY_TYPES=CONTENT_ONLY_TYPES,
        q=q,
        active_statement_id=active_statement_id,
        can_edit=can_edit,
        owner_username=owner_name,
    )

@app.route('/create_topic', methods=['POST'])
@login_required
def create_topic():
    domain_id = request.form.get('domain_id')
    name = request.form.get('name', '').strip()
    description = request.form.get('description', '').strip()
    folder_id = request.form.get('folder_id') or None
    if not domain_id or not name:
        flash('Domain and topic name are required')
        return redirect(request.referrer or url_for('all_domains'))
    conn = get_db()
    # Domains are shared: any logged-in user may create a topic in any domain,
    # so the only check is that the domain exists (no per-user domain gate).
    domain = conn.execute('SELECT id FROM domains WHERE id = ?', (domain_id,)).fetchone()
    if not domain:
        conn.close()
        flash('Domain not found')
        return redirect(url_for('all_domains'))
    # Validate any supplied folder belongs to this domain (topics.domain_id has
    # no declared FK to folders, so ownership is checked here in app code).
    if folder_id:
        folder = conn.execute(
            'SELECT id FROM folders WHERE id = ? AND domain_id = ?',
            (folder_id, domain_id),
        ).fetchone()
        if not folder:
            folder_id = None
    is_public = 1 if request.form.get('is_public') else 0
    conn.execute(
        'INSERT INTO topics (domain_id, folder_id, name, description, user_id, is_public) VALUES (?, ?, ?, ?, ?, ?)',
        (domain_id, folder_id, name, description, g.user['id'], is_public),
    )
    conn.commit()
    topic_id = conn.execute('SELECT last_insert_rowid()').fetchone()[0]
    conn.close()
    return redirect(url_for('topic', topic_id=topic_id))

@app.route('/create_folder', methods=['POST'])
@login_required
def create_folder():
    domain_id = request.form.get('domain_id')
    name = request.form.get('name', '').strip()
    description = request.form.get('description', '').strip()
    parent_id = request.form.get('parent_id') or None
    if not domain_id or not name:
        flash('Domain and folder name are required')
        return redirect(request.referrer or url_for('all_domains'))
    conn = get_db()
    # Domains are shared: any logged-in user may create a folder in any domain,
    # so the only check is that the domain exists (no per-user domain gate).
    domain = conn.execute('SELECT id FROM domains WHERE id = ?', (domain_id,)).fetchone()
    if not domain:
        conn.close()
        flash('Domain not found')
        return redirect(url_for('all_domains'))
    # A parent folder must exist and belong to this domain.
    if parent_id:
        parent = conn.execute(
            'SELECT id FROM folders WHERE id = ? AND domain_id = ?',
            (parent_id, domain_id),
        ).fetchone()
        if not parent:
            parent_id = None
    is_public = 1 if request.form.get('is_public') else 0
    conn.execute(
        'INSERT INTO folders (domain_id, parent_id, name, description, user_id, is_public) VALUES (?, ?, ?, ?, ?, ?)',
        (domain_id, parent_id, name, description, g.user['id'], is_public),
    )
    conn.commit()
    conn.close()
    return redirect(url_for('domain', domain_id=domain_id))

@app.route('/update_folder', methods=['POST'])
@login_required
def update_folder():
    folder_id = request.form.get('folder_id')
    name = request.form.get('name', '').strip()
    description = request.form.get('description', '').strip()
    parent_id = request.form.get('parent_id') or None
    is_public = 1 if request.form.get('is_public') else 0
    if not folder_id or not name:
        flash('Folder name is required')
        return redirect(request.referrer or url_for('all_domains'))
    conn = get_db()
    folder = conn.execute('SELECT f.* FROM folders f WHERE f.id = ?', (folder_id,)).fetchone()
    if not folder or folder['user_id'] != g.user['id']:
        conn.close()
        flash('Folder not found')
        return redirect(request.referrer or url_for('all_domains'))
    # Re-parenting must not create a cycle: the new parent cannot be the folder
    # itself or any of its descendants (walk the candidate parent's chain down,
    # or equivalently check the folder is not within the new parent's subtree).
    if parent_id:
        if int(parent_id) == folder['id']:
            conn.close()
            flash('A folder cannot be its own parent')
            return redirect(url_for('domain', domain_id=folder['domain_id']))
        # A folder may only be re-parented within its own domain.
        new_parent = conn.execute(
            'SELECT id FROM folders WHERE id = ? AND domain_id = ?',
            (parent_id, folder['domain_id']),
        ).fetchone()
        if not new_parent:
            conn.close()
            flash('Parent folder not found in this domain')
            return redirect(url_for('domain', domain_id=folder['domain_id']))
        # Reject if the new parent is a descendant of this folder (cycle guard),
        # using the recursive subtree CTE over child folders.
        cycle = conn.execute('''
            WITH RECURSIVE subtree(id) AS (
                SELECT id FROM folders WHERE id = ?
                UNION ALL
                SELECT f.id FROM folders f JOIN subtree s ON f.parent_id = s.id
            )
            SELECT 1 FROM subtree WHERE id = ?
        ''', (folder['id'], parent_id)).fetchone()
        if cycle:
            conn.close()
            flash('Cannot move a folder into one of its own sub-folders')
            return redirect(url_for('domain', domain_id=folder['domain_id']))
    conn.execute(
        'UPDATE folders SET name = ?, description = ?, parent_id = ?, is_public = ? WHERE id = ?',
        (name, description, parent_id, is_public, folder_id),
    )
    conn.commit()
    conn.close()
    return redirect(url_for('domain', domain_id=folder['domain_id']))

@app.route('/delete_folder/<int:folder_id>', methods=['POST'])
@login_required
def delete_folder(folder_id):
    conn = get_db()
    folder = conn.execute('SELECT f.* FROM folders f WHERE f.id = ?', (folder_id,)).fetchone()
    if not folder or folder['user_id'] != g.user['id']:
        conn.close()
        flash('Folder not found')
        return redirect(request.referrer or url_for('all_domains'))
    domain_id = folder['domain_id']
    # Gather every attachment filename across the whole subtree BEFORE deleting,
    # so the cascade cannot remove the rows that tell us which files to purge.
    subtree_folders, subtree_topics = get_folder_subtree(conn, folder_id)
    if subtree_topics:
        placeholders = ','.join('?' * len(subtree_topics))
        rows = conn.execute('''
            SELECT a.filename FROM attachments a
            JOIN statements s ON a.statement_id = s.id
            WHERE s.topic_id IN (%s) AND a.filename IS NOT NULL
        ''' % placeholders, subtree_topics).fetchall()
    else:
        rows = []
    filenames = [r['filename'] for r in rows if r['filename']]
    # `topics.folder_id` has no declared FK (SQLite cannot attach one to an
    # ALTER-added column), so child topics must be removed explicitly. The folder
    # cascade (parent →
    # child folders, and topics → statements → attachments) still runs; we just
    # delete the subtree topics first, then the folder row last.
    if subtree_topics:
        placeholders = ','.join('?' * len(subtree_topics))
        conn.execute('DELETE FROM topics WHERE id IN (%s)' % placeholders, subtree_topics)
    conn.execute('DELETE FROM folders WHERE id = ?', (folder_id,))
    conn.commit()
    conn.close()
    for fn in filenames:
        _remove_upload(fn)
    return redirect(url_for('domain', domain_id=domain_id))

@app.route('/create_statement', methods=['POST'])
@login_required
def create_statement():
    topic_id = request.form.get('topic_id')
    text = request.form.get('text', '').strip()
    if not topic_id or not text:
        flash('Statement text is required')
        return redirect(request.referrer or url_for('all_domains'))
    conn = get_db()
    topic = require_topic_owner(conn, topic_id)
    if not topic:
        conn.close()
        flash('Topic not found')
        return redirect(request.referrer or url_for('all_domains'))
    conn.execute('INSERT INTO statements (topic_id, text, position) VALUES (?, ?, (SELECT COALESCE(MAX(position), -1) + 1 FROM statements WHERE topic_id = ?))', (topic_id, text, topic_id))
    conn.commit()
    conn.close()
    return redirect(url_for('topic', topic_id=topic_id))

@app.route('/create_attachment', methods=['POST'])
@login_required
def create_attachment():
    # The smart attach modal never asks the user for a type; we infer it from
    # whatever they provided (an uploaded file's extension, a URL, or free
    # text). Title and tags are optional; one of file/content is required.
    statement_id = request.form.get('statement_id')
    title = request.form.get('title', '').strip()
    content = request.form.get('content', '').strip()
    tags = request.form.get('tags', '').strip()
    file = request.files.get('file')

    if not statement_id:
        flash('Statement is required')
        return redirect(request.referrer or url_for('all_domains'))
    conn = get_db()
    stmt = require_statement_owner(conn, statement_id)
    if not stmt:
        conn.close()
        flash('Statement not found')
        return redirect(request.referrer or url_for('all_domains'))

    source_name = (file.filename if file and file.filename else '') or ''
    ext = source_name.rsplit('.', 1)[1].lower() if '.' in source_name else ''

    # Resolve the attachment type from the inputs, in priority order:
    #   1. an uploaded file -> type by extension (image/video/table/document)
    #   2. a bare URL      -> link
    #   3. any other text  -> richtext
    att_type = None
    if file and file.filename:
        att_type = _infer_type_from_filename(file.filename)
    elif _looks_like_url(content):
        att_type = 'link'
    elif content:
        att_type = 'richtext'

    if att_type is None:
        conn.close()
        flash('Add a title with some text, a URL, or upload a file')
        return redirect(request.referrer or url_for('all_domains'))

    # Default the title from the file name when the user left it blank.
    if not title:
        if source_name:
            title = os.path.splitext(os.path.basename(source_name))[0] or 'Untitled'
        elif content:
            title = (content[:60] + '…') if len(content) > 60 else content
        else:
            title = 'Untitled'

    filename = None

    if file and file.filename:
        if _is_dangerous_extension(file.filename):
            conn.close()
            flash('That file type cannot be stored for security reasons')
            return redirect(request.referrer or url_for('all_domains'))

        # A document or richtext upload whose extension we can read as text
        # (txt/md/rtf/doc/docx/odt/pdf) is pulled into native richtext storage
        # instead of being kept as an opaque file download.
        if att_type in ('richtext', 'document') and ext in RICHTEXT_DOC_EXTENSIONS:
            # Pull the document's text into native richtext storage instead of
            # keeping the binary file around as an opaque download.
            tmp = f"{uuid.uuid4().hex}.{ext}"
            os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
            file.save(os.path.join(app.config['UPLOAD_FOLDER'], tmp))
            extracted = _extract_text_from_file(tmp, ext)
            if extracted is not None:
                att_type = 'richtext'
                content = extracted
                filename = None
                os.remove(os.path.join(app.config['UPLOAD_FOLDER'], tmp))
            else:
                # Couldn't read it: fall back to storing the file as a document.
                att_type = 'document'
                filename = tmp
                content = tmp
        elif att_type == 'table':
            # Save the file, then try to convert a spreadsheet into native CSV
            # Univer storage so it opens editable in-browser.
            filename = f"{uuid.uuid4().hex}.{ext}"
            os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
            file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
            csv_text = _table_to_univer_csv(filename, ext)
            if csv_text is not None:
                content = UNIVER_DATA_PREFIX + csv_text
                os.remove(os.path.join(app.config['UPLOAD_FOLDER'], filename))
                filename = None
            else:
                content = filename
        elif att_type in ALLOWED_EXTENSIONS:
            # image/video/document: store the uploaded bytes, point content at it.
            # A document extension we don't have an explicit allowlist entry for
            # is kept as a generic "new file" rather than rejected.
            upload_ext = _upload_extension(att_type, file.filename)
            if upload_ext is None:
                att_type = 'document'
                upload_ext = ext or 'bin'
            filename = f"{uuid.uuid4().hex}.{upload_ext}"
            os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
            file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
            content = filename
        else:
            # Unrecognised upload -> generic "new file" document.
            att_type = 'document'
            upload_ext = ext or 'bin'
            filename = f"{uuid.uuid4().hex}.{upload_ext}"
            os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
            file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
            content = filename
    elif att_type == 'link' and not content:
        conn.close()
        flash('A URL is required for a link asset')
        return redirect(request.referrer or url_for('all_domains'))

    if att_type not in ASSET_TYPES:
        att_type = 'document'

    try:
        conn.execute(
            'INSERT INTO attachments (statement_id, title, type, content, filename, tags, position) '
            'VALUES (?, ?, ?, ?, ?, ?, (SELECT COALESCE(MAX(position), -1) + 1 FROM attachments WHERE statement_id = ?))',
            (statement_id, title, att_type, content, filename, tags, statement_id)
        )
        conn.commit()
        topic_row = conn.execute(
            'SELECT topic_id FROM statements WHERE id = ?', (statement_id,)
        ).fetchone()
    except sqlite3.IntegrityError:
        conn.rollback()
        conn.close()
        flash('Could not save this attachment')
        return redirect(request.referrer or url_for('all_domains'))
    conn.close()
    return redirect_to_topic_referrer(topic_row['topic_id'] if topic_row else None, statement_id)

def _infer_type_from_filename(filename):
    """Best-effort asset type from a dropped file's extension.

    Walks the per-kind allowlists in ALLOWED_EXTENSIONS and returns the first
    matching type. Anything else (an extension we do not recognise) is treated
    as a generic document rather than rejected, so a drop never loses a file --
    only executable/script extensions that would be served as active content are
    refused (see _is_dangerous_extension).
    """
    if '.' not in filename:
        return 'document'
    ext = filename.rsplit('.', 1)[1].lower()
    for att_type, exts in ALLOWED_EXTENSIONS.items():
        if ext in exts:
            return att_type
    return 'document'


# Extensions that must never be stored and served same-origin, because a browser
# could execute them (HTML/SVG as markup, scripts/exes as code). These are the
# only files a drop is allowed to reject; everything else becomes a document.
DANGEROUS_EXTENSIONS = (
    'html', 'htm', 'xhtml', 'svg', 'js', 'mjs', 'php', 'phtml', 'php3',
    'php4', 'php5', 'asp', 'aspx', 'jsp', 'sh', 'bash', 'py', 'pl', 'cgi',
    'exe', 'dll', 'bat', 'cmd', 'com', 'msi', 'scr', 'vbs', 'wsf', 'jar',
    'ps1', 'bin', 'app',
)


def _is_dangerous_extension(filename):
    """True for extensions a browser or OS could execute if served/stored."""
    if '.' not in filename:
        return False
    ext = filename.rsplit('.', 1)[1].lower()
    return ext in DANGEROUS_EXTENSIONS


# Document extensions whose text payload we can pull into native richtext
# storage on upload, instead of keeping the file as an opaque download. Plain
# text/markdown/rtf are read directly; the heavier office formats are parsed
# with the libraries available in this environment.
RICHTEXT_DOC_EXTENSIONS = ('txt', 'md', 'markdown', 'rtf', 'doc', 'docx', 'odt', 'pdf')


def _looks_like_url(value):
    """Loose URL sniff: a scheme, or a host with a recognisable TLD/path."""
    v = (value or '').strip()
    if not v:
        return False
    if re.match(r'^[a-z][a-z0-9+.-]*://', v, re.I):
        return True
    return bool(re.match(r'^(www\.)?[a-z0-9-]+(\.[a-z]{2,})+(/|$)', v, re.I))


def _extract_text_from_file(filename, ext):
    """Best-effort plain-text extraction for a richtext document upload.

    Returns the extracted text, or None when nothing usable could be read (the
    caller then stores the file as an opaque document instead of richtext).
    """
    ext = (ext or '').lower()
    try:
        path = os.path.realpath(os.path.join(app.config['UPLOAD_FOLDER'], filename))
        upload_dir = os.path.realpath(app.config['UPLOAD_FOLDER'])
        if os.path.commonpath([upload_dir, path]) != upload_dir:
            return None
        if ext in ('txt', 'md', 'markdown'):
            with open(path, 'r', encoding='utf-8', errors='replace') as fh:
                return fh.read()
        if ext == 'rtf':
            return _extract_rtf(path)
        if ext == 'docx':
            try:
                import docx
                document = docx.Document(path)
                return '\n'.join(p.text for p in document.paragraphs)
            except Exception:
                return None
        if ext == 'doc':
            try:
                import docx2txt
                return docx2txt.process(path) or None
            except Exception:
                return None
        if ext == 'odt':
            try:
                import xml.etree.ElementTree as ET
                with open(path, 'rb') as fh:
                    root = ET.fromstring(fh.read())
                ns = '{urn:oasis:names:tc:opendocument:xmlns:text:1.0}'
                return '\n'.join(node.text or '' for node in root.iter(ns + 'p'))
            except Exception:
                return None
        if ext == 'pdf':
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(path)
                return '\n'.join((page.extract_text() or '') for page in reader.pages)
            except Exception:
                return None
    except Exception:
        return None
    return None


def _extract_rtf(path):
    """Strip RTF control words, keeping the visible text as plain text."""
    try:
        with open(path, 'r', encoding='utf-8', errors='replace') as fh:
            data = fh.read()
    except OSError:
        return None
    text = re.sub(r'\\[a-z]+\d*\s?', '', data, flags=re.I)
    text = re.sub(r'[{}]', '', text)
    text = re.sub(r'\\\n', '\n', text)
    return text.strip()


def _table_to_univer_csv(filename, ext):
    """Parse a spreadsheet upload into a CSV string for native Univer storage.

    csv/tsv are read directly; xls/xlsx/ods are parsed with openpyxl when the
    first sheet is reachable. Returns the CSV text, or None for binary formats
    we cannot parse (the caller then keeps the file as a table with a
    download-only fallback).
    """
    ext = (ext or '').lower()
    try:
        path = os.path.realpath(os.path.join(app.config['UPLOAD_FOLDER'], filename))
        upload_dir = os.path.realpath(app.config['UPLOAD_FOLDER'])
        if os.path.commonpath([upload_dir, path]) != upload_dir:
            return None
        if ext in ('csv', 'tsv'):
            with open(path, 'r', encoding='utf-8', errors='replace') as fh:
                return fh.read()
        if ext in ('xls', 'xlsx', 'ods'):
            try:
                import openpyxl
                wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
                ws = wb.active
                buf = io.StringIO()
                writer = csv.writer(buf)
                for row in ws.iter_rows(values_only=True):
                    writer.writerow(['' if c is None else c for c in row])
                return buf.getvalue()
            except Exception:
                return None
    except Exception:
        return None
    return None

@app.route('/upload_drop/<int:statement_id>', methods=['POST'])
@login_required
def upload_drop(statement_id):
    """Receive files dropped onto a statement's asset section.

    Used by the drag-and-drop affordance in the evidence pane. The target
    statement is taken from the URL (the currently selected one) rather than a
    form field, so a dropped file can never be reassigned to another statement.
    The asset type is inferred from the file extension, and the title defaults
    to the original filename; both mirror create_attachment's validation.
    """
    conn = get_db()
    statement = require_statement_owner(conn, statement_id)
    if not statement:
        conn.close()
        return {'error': 'Statement not found'}, 404

    results = []
    files = request.files.getlist('file')
    for file in files:
        if not file or not file.filename:
            continue
        source_name = file.filename
        if _is_dangerous_extension(source_name):
            results.append({'filename': source_name, 'error': 'File type is not allowed for security reasons'})
            continue
        # Unknown extensions fall back to the generic document type rather than
        # being rejected, so a drop never silently drops a user's file. The real
        # extension is preserved for those, while recognised types still go
        # through the strict per-type allowlist.
        att_type = _infer_type_from_filename(source_name)
        ext = source_name.rsplit('.', 1)[1].lower() if '.' in source_name else ''
        title = os.path.splitext(os.path.basename(source_name))[0] or 'Untitled'

        filename = None
        content = ''
        if file and ext in RICHTEXT_DOC_EXTENSIONS and att_type in ('document', 'richtext'):
            # Document dropped: extract its text into native richtext storage.
            tmp = f"{uuid.uuid4().hex}.{ext}"
            os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
            file.save(os.path.join(app.config['UPLOAD_FOLDER'], tmp))
            extracted = _extract_text_from_file(tmp, ext)
            if extracted is not None:
                att_type = 'richtext'
                content = extracted
                os.remove(os.path.join(app.config['UPLOAD_FOLDER'], tmp))
            else:
                att_type = 'document'
                filename = tmp
                content = tmp
        elif att_type == 'table':
            upload_ext = ext or 'csv'
            filename = f"{uuid.uuid4().hex}.{upload_ext}"
            os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
            file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
            csv_text = _table_to_univer_csv(filename, upload_ext)
            if csv_text is not None:
                content = UNIVER_DATA_PREFIX + csv_text
                os.remove(os.path.join(app.config['UPLOAD_FOLDER'], filename))
                filename = None
            else:
                content = filename
        else:
            upload_ext = _upload_extension(att_type, source_name)
            if upload_ext is None:
                if att_type == 'document':
                    upload_ext = ext
                else:
                    results.append({'filename': source_name, 'error': 'Unsupported file type'})
                    continue
            filename = f"{uuid.uuid4().hex}.{upload_ext}"
            os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
            file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
            content = filename

        if att_type not in ASSET_TYPES:
            att_type = 'document'
        try:
            cursor = conn.execute(
                'INSERT INTO attachments (statement_id, title, type, content, filename, tags, position) '
                'VALUES (?, ?, ?, ?, ?, ?, (SELECT COALESCE(MAX(position), -1) + 1 FROM attachments WHERE statement_id = ?))',
                (statement_id, title, att_type, content, filename, '', statement_id)
            )
            conn.commit()
            new_id = cursor.lastrowid
            results.append({
                'filename': source_name,
                'id': new_id,
                'attachment': {
                    'id': new_id,
                    'statement_id': statement_id,
                    'title': title,
                    'type': att_type,
                    'kind': asset_kind(att_type, filename),
                    'preview': '',
                    'filename': filename,
                    'tags': '',
                },
            })
        except sqlite3.IntegrityError:
            conn.rollback()
            results.append({'filename': source_name, 'error': 'Could not save attachment'})

    conn.close()
    return {'results': results}

@app.route('/update_statement', methods=['POST'])
@login_required
def update_statement():
    statement_id = request.form.get('statement_id')
    text = request.form.get('text', '').strip()
    if not statement_id or not text:
        flash('Statement text is required')
        return redirect(request.referrer or url_for('all_domains'))
    conn = get_db()
    stmt = require_statement_owner(conn, statement_id)
    if not stmt:
        conn.close()
        flash('Statement not found')
        return redirect(request.referrer or url_for('all_domains'))
    conn.execute('UPDATE statements SET text = ? WHERE id = ?', (text, statement_id))
    conn.commit()
    conn.close()
    return redirect(request.referrer or url_for('all_domains'))

@app.route('/reorder_statements/<int:topic_id>', methods=['POST'])
@login_required
def reorder_statements(topic_id):
    # Receives the full ordered list of statement ids for the topic; each is
    # assigned a sequential position so the new order is authoritative.
    ordered = request.form.getlist('order')
    if not ordered:
        return redirect(request.referrer or url_for('all_domains'))
    conn = get_db()
    topic = require_topic_owner(conn, topic_id)
    if not topic:
        conn.close()
        flash('Topic not found')
        return redirect(request.referrer or url_for('all_domains'))
    try:
        conn.executemany(
            'UPDATE statements SET position = ? WHERE id = ? AND topic_id = ?',
            [(idx, int(sid), topic_id) for idx, sid in enumerate(ordered)],
        )
        conn.commit()
    except (sqlite3.IntegrityError, ValueError):
        conn.rollback()
        conn.close()
        flash('Could not reorder statements')
        return redirect(request.referrer or url_for('all_domains'))
    conn.close()
    return redirect(request.referrer or url_for('all_domains'))

@app.route('/reorder_attachments/<int:statement_id>', methods=['POST'])
@login_required
def reorder_attachments(statement_id):
    # Like reorder_statements but scoped to one statement's assets.
    ordered = request.form.getlist('order')
    if not ordered:
        return redirect(request.referrer or url_for('all_domains'))
    conn = get_db()
    stmt = require_statement_owner(conn, statement_id)
    if not stmt:
        conn.close()
        flash('Statement not found')
        return redirect(request.referrer or url_for('all_domains'))
    try:
        conn.executemany(
            'UPDATE attachments SET position = ? WHERE id = ? AND statement_id = ?',
            [(idx, int(aid), statement_id) for idx, aid in enumerate(ordered)],
        )
        conn.commit()
    except (sqlite3.IntegrityError, ValueError):
        conn.rollback()
        conn.close()
        flash('Could not reorder assets')
        return redirect(request.referrer or url_for('all_domains'))
    conn.close()
    return redirect(request.referrer or url_for('all_domains'))

@app.route('/delete_statement/<int:statement_id>', methods=['POST'])
@login_required
def delete_statement(statement_id):
    conn = get_db()
    stmt = require_statement_owner(conn, statement_id)
    if not stmt:
        conn.close()
        flash('Statement not found')
        return redirect(request.referrer or url_for('all_domains'))
    # Collect files before the cascade removes the rows that reference them.
    rows = conn.execute(
        'SELECT filename FROM attachments WHERE statement_id = ? AND filename IS NOT NULL',
        (statement_id,)
    ).fetchall()
    topics = conn.execute(
        'SELECT id, topic_id FROM statements WHERE topic_id = (SELECT topic_id FROM statements WHERE id = ?) ORDER BY position ASC, created_at ASC',
        (statement_id,)
    ).fetchall()
    conn.execute('DELETE FROM statements WHERE id = ?', (statement_id,))
    conn.commit()
    topic_id = topics[0]['topic_id'] if topics else None
    # Keep the user on the next sibling statement, falling back to the first
    # remaining one so the view doesn't jump back to the top of the list.
    next_stmt = None
    for t in topics:
        if t['id'] == statement_id:
            continue
        next_stmt = t['id']
        break
    if next_stmt is None and len(topics) > 1:
        next_stmt = topics[0]['id']
    conn.close()
    for row in rows:
        _remove_upload(row['filename'])
    return redirect_to_topic_referrer(topic_id, next_stmt)

@app.route('/update_attachment', methods=['POST'])
@login_required
def update_attachment():
    attachment_id = request.form.get('attachment_id')
    title = request.form.get('title', '').strip()
    att_type = request.form.get('type', '')
    content = request.form.get('content', '').strip()
    tags = request.form.get('tags', '').strip()
    file = request.files.get('file')

    if not attachment_id or not title:
        flash('Attachment title is required')
        return redirect(request.referrer or url_for('all_domains'))

    conn = get_db()
    existing = conn.execute('SELECT a.*, t.user_id FROM attachments a JOIN statements s ON a.statement_id = s.id JOIN topics t ON s.topic_id = t.id WHERE a.id = ?', (attachment_id,)).fetchone()
    if not existing or existing['user_id'] != g.user['id']:
        conn.close()
        flash('Attachment not found')
        return redirect(request.referrer or url_for('all_domains'))

    if att_type not in ASSET_TYPES:
        att_type = existing['type']

    filename = existing['filename']
    old_filename = existing['filename']

    if att_type not in CONTENT_ONLY_TYPES and file and file.filename:
        ext = _upload_extension(att_type, file.filename)
        if ext is None:
            conn.close()
            flash('That file type is not allowed for this asset type')
            return redirect(request.referrer or url_for('all_domains'))
        filename = f"{uuid.uuid4().hex}.{ext}"
        os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
        file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
        content = filename
    elif att_type in CONTENT_ONLY_TYPES:
        # These types are content-only; drop any previously attached file.
        filename = None
        if not content:
            conn.close()
            flash('Content is required for this asset type')
            return redirect(request.referrer or url_for('all_domains'))
    elif old_filename:
        # File-backed type keeping its existing file. The edit form still
        # submits the (hidden) URL field, so never trust it to replace the
        # stored value; only fall back to the filename when content is empty
        # or the two already agree, so a legitimate label/URL is preserved.
        filename = old_filename
        if not content or content == existing['content']:
            content = existing['content'] or old_filename
    elif content:
        # File-backed type pointing at a URL rather than an upload (e.g. a
        # YouTube video). All three renderers support this filename-less state.
        filename = None
    else:
        # Neither a file, an existing file, nor a URL: this would persist a
        # corrupt row that renders as /uploads/<empty>.
        conn.close()
        flash('A file or URL is required for this asset type')
        return redirect(request.referrer or url_for('all_domains'))

    try:
        conn.execute(
            'UPDATE attachments SET title = ?, type = ?, content = ?, filename = ?, tags = ? WHERE id = ?',
            (title, att_type, content, filename, tags, attachment_id)
        )
        conn.commit()
    except sqlite3.IntegrityError:
        conn.rollback()
        conn.close()
        flash('Could not save this attachment')
        return redirect(request.referrer or url_for('all_domains'))
    conn.close()

    if old_filename and old_filename != filename:
        _remove_upload(old_filename)

    topic_row = None
    if existing:
        tconn = get_db()
        topic_row = tconn.execute(
            'SELECT topic_id FROM statements WHERE id = ?', (existing['statement_id'],)
        ).fetchone()
        tconn.close()
    return redirect_to_topic_referrer(topic_row['topic_id'] if topic_row else None, existing['statement_id'] if existing else None)


@app.route('/delete_attachment/<int:attachment_id>', methods=['POST'])
@login_required
def delete_attachment(attachment_id):
    conn = get_db()
    row = conn.execute(
        'SELECT a.filename, a.statement_id, t.user_id FROM attachments a '
        'JOIN statements s ON a.statement_id = s.id '
        'JOIN topics t ON s.topic_id = t.id WHERE a.id = ?', (attachment_id,)
    ).fetchone()
    if not row or row['user_id'] != g.user['id']:
        conn.close()
        flash('Attachment not found')
        return redirect(request.referrer or url_for('all_domains'))
    statement_id = row['statement_id'] if row else None
    conn.execute('DELETE FROM attachments WHERE id = ?', (attachment_id,))
    conn.commit()
    topic_row = conn.execute(
        'SELECT topic_id FROM statements WHERE id = ?', (statement_id,)
    ).fetchone() if statement_id else None
    conn.close()
    if row and row['filename']:
        _remove_upload(row['filename'])
    return redirect_to_topic_referrer(topic_row['topic_id'] if topic_row else None, statement_id)


@app.route('/attachment/<int:attachment_id>')
def attachment(attachment_id):
    # Full content is fetched on demand so the initial topic page only ships
    # previews, not every richtext body. The modal and edit form call this.
    conn = get_db()
    row = conn.execute(
        'SELECT id, statement_id, title, type, content, filename, tags FROM attachments WHERE id = ?',
        (attachment_id,)
    ).fetchone()
    conn.close()
    if not row:
        return {'error': 'not found'}, 404
    payload = {
        'id': row['id'],
        'statement_id': row['statement_id'],
        'title': row['title'],
        'type': row['type'],
        'kind': asset_kind(row['type'], row['filename']),
        'content': row['content'],
        'filename': row['filename'] or '',
        'tags': row['tags'] or '',
    }
    return payload


def _table_to_rows(filename, content):
    """Return a list-of-rows (list of lists of strings) for a table asset.

    `content` is either an uploaded filename (read from disk) or, after an
    in-browser Univer edit, a `univer:`-prefixed CSV payload stored directly in
    the column. CSV/TSV are parsed with the stdlib; the binary spreadsheet
    formats (xls/xlsx/ods) are returned as None because we cannot parse them
    without an extra dependency, in which case the client falls back to a
    download link and a read-only view is unavailable.
    """
    raw = content
    if content and not content.startswith(UNIVER_DATA_PREFIX) and filename:
        # Uploaded file: read its bytes from the upload folder.
        try:
            path = os.path.realpath(os.path.join(app.config['UPLOAD_FOLDER'], filename))
            upload_dir = os.path.realpath(app.config['UPLOAD_FOLDER'])
            if os.path.commonpath([upload_dir, path]) == upload_dir:
                with open(path, 'rb') as fh:
                    raw = fh.read().decode('utf-8', errors='replace')
        except OSError:
            raw = content

    if raw and raw.startswith(UNIVER_DATA_PREFIX):
        raw = raw[len(UNIVER_DATA_PREFIX):]

    if not raw:
        return None

    ext = filename.rsplit('.', 1)[1].lower() if filename and '.' in filename else 'csv'
    # Binary spreadsheet formats cannot be parsed with the stdlib, so signal
    # "unsupported" and let the client keep the download-only fallback. Parsing
    # raw bytes as CSV would corrupt the file into a single bogus cell.
    if ext in ('xls', 'xlsx', 'ods'):
        return None

    delimiter = '\t' if ext == 'tsv' else ','
    try:
        reader = csv.reader(io.StringIO(raw), delimiter=delimiter)
        return [list(row) for row in reader]
    except (csv.Error, ValueError):
        return None


@app.route('/attachment/<int:attachment_id>/table')
def attachment_table(attachment_id):
    """Return a table asset as JSON rows for the Univer editor.

    Uploaded CSV/TSV files and previously in-browser-edited tables resolve to a
    row matrix. Binary spreadsheet formats (xls/xlsx/ods) cannot be parsed by
    the stdlib, so they come back with `supported: false` and the client keeps
    the download-only behaviour.
    """
    conn = get_db()
    row = conn.execute(
        'SELECT id, type, content, filename FROM attachments WHERE id = ?',
        (attachment_id,)
    ).fetchone()
    conn.close()
    if not row:
        return jsonify({'error': 'not found'}), 404
    if asset_kind(row['type'], row['filename']) != 'table':
        return jsonify({'error': 'not a table'}), 400

    rows = _table_to_rows(row['filename'], row['content'])
    if rows is None:
        return jsonify({'supported': False, 'filename': row['filename'] or ''})
    return jsonify({'supported': True, 'rows': rows, 'filename': row['filename'] or ''})


@app.route('/api/tree')
def api_tree():
    """Return the full domain→folder→topic→statement tree for the extension picker."""
    conn = get_db()
    # The picker is an authenticated owner tool, so no public-only filtering.
    topic_vis = ''
    domains = conn.execute('SELECT id, name, description FROM domains ORDER BY name').fetchall()
    result = {'domains': []}
    for d in domains:
        domain_id = d['id']
        # Folders for this domain
        folders = conn.execute(
            'SELECT id, parent_id, name, description FROM folders WHERE domain_id = ? ORDER BY name', (domain_id,)
        ).fetchall()
        folder_tree = build_folder_tree(conn, folders, topic_vis)
        # Topics with no folder (loose topics)
        loose_topics = conn.execute('''
            SELECT t.id, t.name, t.description,
                   COUNT(DISTINCT s.id) as statement_count
            FROM topics t
            LEFT JOIN statements s ON t.id = s.topic_id
            WHERE t.domain_id = ? AND t.folder_id IS NULL
            GROUP BY t.id
            ORDER BY t.created_at DESC
        ''', (domain_id,)).fetchall()
        loose_topics_list = []
        for t in loose_topics:
            statements = conn.execute(
                'SELECT id, text FROM statements WHERE topic_id = ? ORDER BY position ASC, created_at ASC', (t['id'],)
            ).fetchall()
            loose_topics_list.append({
                'id': t['id'],
                'name': t['name'],
                'description': t['description'],
                'statements': [{'id': s['id'], 'text': s['text']} for s in statements],
            })
        # Attach loose_topics to folder_tree root level
        for node in folder_tree:
            node['loose_topics'] = []
        result['domains'].append({
            'id': domain_id,
            'name': d['name'],
            'description': d['description'],
            'folders': folder_tree,
            'loose_topics': loose_topics_list,
        })
    conn.close()
    return jsonify(result)


# nh3 allowlist for CKEditor 5 compatibility
CKEDITOR_TAGS = {
    'p', 'br', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6',
    'strong', 'em', 'u', 's', 'code', 'pre', 'blockquote',
    'ul', 'ol', 'li', 'a', 'img', 'hr', 'figure', 'figcaption',
    'table', 'thead', 'tbody', 'tr', 'th', 'td',
    'span', 'div',
}
CKEDITOR_ATTRS = {
    '*': {'class', 'style', 'id'},
    'a': {'href', 'target'},
    'img': {'src', 'alt', 'width', 'height', 'loading'},
    'th': {'scope', 'colspan', 'rowspan'},
    'td': {'colspan', 'rowspan'},
}

SAFE_URL_SCHEMES = frozenset(('http', 'https', 'data'))


def _sanitize_html(html):
    """Sanitize HTML with nh3 using CKEditor-compatible allowlist."""
    return nh3.clean(
        html,
        tags=CKEDITOR_TAGS,
        attributes=CKEDITOR_ATTRS,
        url_schemes=SAFE_URL_SCHEMES,
    )


def _fetch_and_save_image(url, session, upload_folder, seen_hashes, max_size=20 * 1024 * 1024, timeout=10):
    """Fetch an image, validate, deduplicate by SHA256, save to uploads. Returns (local_path, filename) or (None, None)."""
    try:
        resp = session.get(url, timeout=timeout, headers={'Accept': 'image/*'}, stream=True)
        if resp.status_code != 200:
            return None, None
        content_type = resp.headers.get('Content-Type', '').lower()
        if not content_type.startswith('image/'):
            return None, None
        # Read with size limit
        content = b''
        for chunk in resp.iter_content(chunk_size=8192):
            content += chunk
            if len(content) > max_size:
                return None, None
        # Hash for deduplication
        file_hash = hashlib.sha256(content).hexdigest()
        if file_hash in seen_hashes:
            return seen_hashes[file_hash], seen_hashes[file_hash]
        # Determine extension from MIME type or URL
        ext_map = {
            'image/jpeg': 'jpg',
            'image/png': 'png',
            'image/gif': 'gif',
            'image/webp': 'webp',
            'image/bmp': 'bmp',
            'image/avif': 'avif',
        }
        ext = ext_map.get(content_type)
        if not ext:
            # Fallback to URL extension
            url_path = url.split('?')[0]
            if '.' in url_path:
                possible_ext = url_path.rsplit('.', 1)[1].lower()
                if possible_ext in ('jpg', 'jpeg', 'png', 'gif', 'webp', 'bmp', 'avif'):
                    ext = 'jpg' if possible_ext == 'jpeg' else possible_ext
        if not ext:
            ext = 'jpg'
        filename = f"{uuid.uuid4().hex}.{ext}"
        filepath = os.path.join(upload_folder, filename)
        os.makedirs(upload_folder, exist_ok=True)
        with open(filepath, 'wb') as f:
            f.write(content)
        seen_hashes[file_hash] = (f"/uploads/{filename}", filename)
        return f"/uploads/{filename}", filename
    except Exception:
        return None, None


def _rewrite_image_srcs(html, session, upload_folder, seen_hashes):
    """Find all <img src="..."> in HTML, fetch/rehost each, rewrite src to local path."""
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(html, 'html.parser')
    for img in soup.find_all('img'):
        src = img.get('src')
        if not src:
            continue
        # Skip data URIs and already-local uploads
        if src.startswith('data:') or src.startswith('/uploads/'):
            continue
        # Make absolute if relative
        if src.startswith('//'):
            src = 'https:' + src
        elif src.startswith('/'):
            # We don't have the base URL here; skip relative paths
            continue
        local_path, filename = _fetch_and_save_image(src, session, upload_folder, seen_hashes)
        if local_path:
            img['src'] = local_path
    return str(soup)


def _build_capture(title, url, html):
    """Run the shared webpage-capture pipeline: readability → sanitize →
    rehost images → compose a richtext body. Returns (source_title, composed_html).
    Used by both the extension's /api/capture (which inserts) and the on-site
    scrape tab's /api/scrape_preview (which returns the fields for review)."""
    # 1. Readability extraction
    try:
        doc = readability.Document(html)
        article_html = doc.summary()
        extracted_title = doc.title() or ''
    except Exception as e:
        raise RuntimeError(f'Readability failed: {e}')

    # 2. Sanitize article HTML
    clean_html = _sanitize_html(article_html)

    # 3. Image harvest & rehost
    upload_folder = app.config['UPLOAD_FOLDER']
    session = requests.Session()
    seen_hashes = {}
    final_html = _rewrite_image_srcs(clean_html, session, upload_folder, seen_hashes)

    # 4. Compose richtext body
    source_title = title or extracted_title or url
    source_link = f'<p><strong>Source:</strong> <a href="{nh3.clean(url, tags=set(), attributes={}, url_schemes=SAFE_URL_SCHEMES)}">{nh3.clean(source_title, tags=set(), attributes={})}</a></p>'
    # Excerpt: first paragraph or readability excerpt
    excerpt_html = ''
    soup_excerpt = BeautifulSoup(final_html, 'html.parser')
    first_p = soup_excerpt.find('p')
    if first_p and first_p.get_text(strip=True):
        excerpt_html = f'<p>{first_p.get_text(strip=True)[:300]}</p>'
    hr = '<hr>'
    composed = f'{source_link}{excerpt_html}{hr}{final_html}'
    return source_title, composed


@app.route('/api/scrape_preview', methods=['POST'])
def api_scrape_preview():
    """On-site scrape tab: fetch a URL (or use supplied html), run the capture
    pipeline, and return the cleaned title + richtext content WITHOUT inserting.
    The client populates the modal fields so the user can review before saving."""
    data = request.get_json(silent=True) or {}
    title = (data.get('title') or '').strip()
    url = (data.get('url') or '').strip()
    html = data.get('html') or ''

    if not url:
        return jsonify({'ok': False, 'error': 'url is required'}), 400

    # The on-site tab only has the URL; fetch it server-side.
    if not html:
        try:
            resp = requests.get(url, timeout=15, headers={'User-Agent': 'Mozilla/5.0 (compatible; CompendiumCapture/1.0)'})
            resp.raise_for_status()
            html = resp.text
        except Exception as e:
            return jsonify({'ok': False, 'error': f'Could not fetch URL: {e}'}), 502

    try:
        source_title, composed = _build_capture(title, url, html)
    except RuntimeError as e:
        return jsonify({'ok': False, 'error': str(e)}), 500

    return jsonify({'ok': True, 'title': source_title, 'content': composed, 'source_url': url})


@app.route('/api/capture', methods=['POST'])
@login_required
def api_capture():
    """Capture a webpage: readability → sanitize → rehost images → save as richtext attachment."""
    data = request.get_json(silent=True) or {}
    statement_id = data.get('statement_id')
    title = (data.get('title') or '').strip()
    url = (data.get('url') or '').strip()
    html = data.get('html') or ''
    tags = (data.get('tags') or '').strip()

    if not statement_id or not url:
        return jsonify({'ok': False, 'error': 'statement_id and url are required'}), 400

    conn = get_db()
    stmt = require_statement_owner(conn, statement_id)
    if not stmt:
        conn.close()
        return jsonify({'ok': False, 'error': 'statement not found'}), 404

    # The browser extension supplies the page's raw HTML; when called from the
    # on-site scrape tab we only have the URL, so fetch it server-side.
    if not html:
        try:
            resp = requests.get(url, timeout=15, headers={'User-Agent': 'Mozilla/5.0 (compatible; CompendiumCapture/1.0)'})
            resp.raise_for_status()
            html = resp.text
        except Exception as e:
            return jsonify({'ok': False, 'error': f'Could not fetch URL: {e}'}), 502

    source_title, composed = _build_capture(title, url, html)
    # 5. Insert as richtext attachment with source_url
    try:
        conn.execute(
            'INSERT INTO attachments (statement_id, title, type, content, filename, tags, source_url, position) '
            'VALUES (?, ?, ?, ?, ?, ?, ?, (SELECT COALESCE(MAX(position), -1) + 1 FROM attachments WHERE statement_id = ?))',
            (statement_id, source_title, 'richtext', composed, None, tags, url, statement_id)
        )
        conn.commit()
        attachment_id = conn.execute('SELECT last_insert_rowid()').fetchone()[0]
    except sqlite3.IntegrityError:
        conn.rollback()
        conn.close()
        return jsonify({'ok': False, 'error': 'Could not save attachment'}), 500
    conn.close()

    return jsonify({'ok': True, 'attachment_id': attachment_id})


@app.route('/save_table/<int:attachment_id>', methods=['POST'])
@login_required
def save_table(attachment_id):
    """Persist an in-browser Univer edit back to the attachment.

    The client exports the sheet to CSV and POSTs it as `csv`. It is stored in
    the `content` column under the `univer:` prefix so it round-trips without a
    binary file, and `filename` is cleared so the row is no longer treated as a
    file-backed asset. A stored file is removed on disk once superseded.
    """
    conn = get_db()
    existing = conn.execute('SELECT a.*, t.user_id FROM attachments a JOIN statements s ON a.statement_id = s.id JOIN topics t ON s.topic_id = t.id WHERE a.id = ?', (attachment_id,)).fetchone()
    if not existing or existing['user_id'] != g.user['id']:
        conn.close()
        return jsonify({'error': 'not found'}), 404

    csv_text = request.form.get('csv', '')
    if not csv_text.strip():
        conn.close()
        return jsonify({'error': 'empty table'}), 400

    old_filename = existing['filename']
    try:
        conn.execute(
            'UPDATE attachments SET content = ?, filename = ? WHERE id = ?',
            (UNIVER_DATA_PREFIX + csv_text, None, attachment_id)
        )
        conn.commit()
    except sqlite3.IntegrityError:
        conn.rollback()
        conn.close()
        return jsonify({'error': 'could not save'}), 500
    conn.close()

    if old_filename:
        _remove_upload(old_filename)
    return jsonify({'ok': True})


def topic_ids_matching(conn, topic_ids, q):
    """Subset of topic_ids whose name/desc, any statement text, or any
    attachment title matches q (case-insensitive LIKE). All params bound."""
    pat = f"%{q}%"
    ids = list(topic_ids)
    if not ids:
        return set()
    ph = ",".join("?" * len(ids))
    matched = set()
    matched |= {r["id"] for r in conn.execute(
        f"SELECT id FROM topics WHERE id IN ({ph}) "
        f"AND (LOWER(name) LIKE LOWER(?) OR LOWER(COALESCE(description,'')) LIKE LOWER(?))",
        ids + [pat, pat])}
    matched |= {r["topic_id"] for r in conn.execute(
        f"SELECT DISTINCT topic_id FROM statements WHERE topic_id IN ({ph}) AND LOWER(text) LIKE LOWER(?)",
        ids + [pat])}
    matched |= {r["topic_id"] for r in conn.execute(
        f"SELECT DISTINCT s.topic_id FROM attachments a JOIN statements s ON a.statement_id=s.id "
        f"WHERE s.topic_id IN ({ph}) AND LOWER(COALESCE(a.title,'')) LIKE LOWER(?)",
        ids + [pat])}
    return matched


def get_folder_subtree(conn, folder_id):
    """Return the full subtree rooted at `folder_id` as `(folder_ids, topic_ids)`.

    Uses a single recursive CTE to walk child folders, then one more CTE to
    collect every topic that hangs anywhere under that subtree (including
    topics attached directly to the root folder). The result lets a caller
    compute aggregate counts or gather attachment filenames across the whole
    tree without hitting the database once per node.
    """
    folder_rows = conn.execute('''
        WITH RECURSIVE subtree(id) AS (
            SELECT id FROM folders WHERE id = ?
            UNION ALL
            SELECT f.id FROM folders f JOIN subtree s ON f.parent_id = s.id
        )
        SELECT id FROM subtree
    ''', (folder_id,)).fetchall()
    folder_ids = [row['id'] for row in folder_rows]

    topic_rows = conn.execute('''
        WITH RECURSIVE subtree(id) AS (
            SELECT id FROM folders WHERE id = ?
            UNION ALL
            SELECT f.id FROM folders f JOIN subtree s ON f.parent_id = s.id
        )
        SELECT id FROM topics WHERE folder_id IN (SELECT id FROM subtree)
    ''', (folder_id,)).fetchall()
    topic_ids = [row['id'] for row in topic_rows]

    return folder_ids, topic_ids


def build_folder_tree(conn, folders, topic_vis='', topic_params=None):
    """Nest a flat folder list into a tree and attach topic + aggregate counts.

    `folders` is any iterable of folder rows for a single domain. Each returned
    node is a dict with the folder's columns plus `children`, `topics`, and
    rolled-up `statement_count` / `attachment_count` over the whole subtree.

    `topic_vis` is an extra SQL predicate (e.g. 'AND t.is_public = 1') appended
    to the per-folder topic query so non-owners only ever see public topics.
    `topic_params` carries any bound parameters used inside `topic_vis`.
    """
    if topic_params is None:
        topic_params = []
    by_id = {}
    for f in folders:
        by_id[f['id']] = {
            'id': f['id'],
            'parent_id': f['parent_id'],
            'name': f['name'],
            'description': f['description'],
            'is_public': f['is_public'],
            'children': [],
            'topics': [],
            'statement_count': 0,
            'attachment_count': 0,
        }

    roots = []
    for node in by_id.values():
        if node['parent_id'] and node['parent_id'] in by_id:
            by_id[node['parent_id']]['children'].append(node)
        else:
            roots.append(node)

    # Aggregate per folder: its own topics then roll children up the tree.
    for node in by_id.values():
        rows = conn.execute('''
            SELECT t.id, t.name, t.description,
                   COUNT(DISTINCT s.id) as statement_count,
                   COUNT(DISTINCT a.id) as attachment_count
            FROM topics t
            LEFT JOIN statements s ON t.id = s.topic_id
            LEFT JOIN attachments a ON s.id = a.statement_id
            WHERE t.folder_id = ? %s
            GROUP BY t.id
            ORDER BY t.created_at DESC
        ''' % topic_vis, (node['id'],) + tuple(topic_params)).fetchall()
        node['topics'] = [dict(r) for r in rows]
        node['statement_count'] += sum(r['statement_count'] for r in rows)
        node['attachment_count'] += sum(r['attachment_count'] for r in rows)

    def roll_up(node):
        for child in node['children']:
            roll_up(child)
            node['statement_count'] += child['statement_count']
            node['attachment_count'] += child['attachment_count']

    for root in roots:
        roll_up(root)

    return roots


def _remove_upload(filename):
    """Delete an uploaded file from disk, ignoring anything outside the upload dir."""
    if not filename:
        return
    upload_dir = os.path.realpath(app.config['UPLOAD_FOLDER'])
    path = os.path.realpath(os.path.join(upload_dir, filename))
    if os.path.commonpath([upload_dir, path]) != upload_dir:
        return
    try:
        os.remove(path)
    except OSError:
        pass


def _copy_attachment(conn, src_att, new_statement_id):
    """Insert a deep copy of `src_att` for `new_statement_id`.

    File-backed attachments get a fresh uuid filename and the physical file is
    copied into UPLOAD_FOLDER; content-only types keep their body verbatim.
    Returns the new attachment row id.
    """
    att_type = src_att['type']
    filename = src_att['filename']
    content = src_att['content']
    if filename and att_type not in CONTENT_ONLY_TYPES:
        ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
        new_filename = f"{uuid.uuid4().hex}.{ext}" if ext else f"{uuid.uuid4().hex}"
        src_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        dst_path = os.path.join(app.config['UPLOAD_FOLDER'], new_filename)
        # Only copy when the source file actually exists; otherwise point the
        # copy at the original name so it still renders rather than 500-ing.
        if os.path.exists(src_path):
            shutil.copy(src_path, dst_path)
            content = new_filename
            filename = new_filename
    cursor = conn.execute(
        'INSERT INTO attachments (statement_id, title, type, content, filename, position) '
        'VALUES (?, ?, ?, ?, ?, ?)',
        (new_statement_id, src_att['title'], att_type, content, filename, src_att['position']),
    )
    return cursor.lastrowid


def _copy_topic(conn, src_topic, new_folder_id=None):
    """Deep-copy `src_topic` (statements + attachments + files) into `new_folder_id`.

    If `new_folder_id` is None the copy keeps the original's folder_id (used by
    duplicate_topic). Returns the new topic id.
    """
    folder_id = src_topic['folder_id'] if new_folder_id is None else new_folder_id
    cursor = conn.execute(
        "INSERT INTO topics (domain_id, folder_id, name, description, user_id, created_at) "
        "VALUES (?, ?, 'Copy of ' || ?, ?, ?, ?)",
        (src_topic['domain_id'], folder_id, src_topic['name'],
         src_topic['description'], src_topic['user_id'], src_topic['created_at']),
    )
    new_topic_id = cursor.lastrowid
    src_statements = conn.execute(
        'SELECT * FROM statements WHERE topic_id = ? ORDER BY position ASC, created_at ASC',
        (src_topic['id'],),
    ).fetchall()
    for src_stmt in src_statements:
        cursor = conn.execute(
            'INSERT INTO statements (topic_id, text, position) VALUES (?, ?, ?)',
            (new_topic_id, src_stmt['text'], src_stmt['position']),
        )
        new_statement_id = cursor.lastrowid
        src_attachments = conn.execute(
            'SELECT * FROM attachments WHERE statement_id = ? ORDER BY position ASC, created_at DESC',
            (src_stmt['id'],),
        ).fetchall()
        for src_att in src_attachments:
            _copy_attachment(conn, src_att, new_statement_id)
    return new_topic_id


@app.route('/move_topic', methods=['POST'])
@login_required
def move_topic():
    topic_id = request.form.get('topic_id')
    folder_id = request.form.get('folder_id') or None
    conn = get_db()
    topic = require_topic_owner(conn, topic_id)
    if not topic:
        conn.close()
        flash('Topic not found')
        return redirect(request.referrer or url_for('all_domains'))
    domain_id = topic['domain_id']
    # Validate the target folder belongs to this domain (topics.folder_id has no
    # declared FK, so ownership is enforced in app code like create_topic does).
    if folder_id:
        folder = conn.execute(
            'SELECT id FROM folders WHERE id = ? AND domain_id = ?',
            (folder_id, domain_id),
        ).fetchone()
        if not folder:
            conn.close()
            flash('Folder not found in this domain')
            return redirect(url_for('domain', domain_id=domain_id))
    conn.execute('UPDATE topics SET folder_id = ? WHERE id = ?', (folder_id, topic_id))
    conn.commit()
    conn.close()
    return redirect(url_for('domain', domain_id=domain_id))


@app.route('/move_statement/<int:statement_id>', methods=['POST'])
@login_required
def move_statement(statement_id):
    """Re-parent a statement under a *different topic* in the SAME domain.

    Semantics (documented in README "Move semantics"):
      - A statement lives directly under a topic (statements.topic_id).
      - "Move" here means changing the statement's topic, i.e. moving it
        sideways/down to a peer topic. It never targets a folder: a folder
        sits ABOVE topics in the hierarchy (domains › folders › topics), so a
        statement can never be filed directly under a folder.
      - The target topic must belong to the statement's current domain; cross-
        domain moves are rejected because a topic is domain-scoped.
    """
    to_topic_id = request.form.get('to_topic_id')
    conn = get_db()
    stmt = require_statement_owner(conn, statement_id)
    if not stmt:
        conn.close()
        flash('Statement not found')
        return redirect(request.referrer or url_for('all_domains'))
    topic = conn.execute('SELECT * FROM topics WHERE id = ?', (to_topic_id,)).fetchone()
    if not topic:
        conn.close()
        flash('Target topic not found')
        return redirect_to_topic_referrer(stmt['topic_id'], statement_id)
    # Domain guard: the statement's topic and the target topic must share a domain.
    if stmt['topic_id'] == to_topic_id:
        conn.close()
        return redirect_to_topic_referrer(stmt['topic_id'], statement_id)
    current_topic = conn.execute('SELECT domain_id FROM topics WHERE id = ?', (stmt['topic_id'],)).fetchone()
    if not current_topic or current_topic['domain_id'] != topic['domain_id']:
        conn.close()
        flash('A statement can only be moved to a topic in the same domain')
        return redirect_to_topic_referrer(stmt['topic_id'], statement_id)
    # Append at the end of the destination topic's statement list.
    max_pos = conn.execute(
        'SELECT MAX(position) AS m FROM statements WHERE topic_id = ?', (to_topic_id,)
    ).fetchone()['m']
    next_pos = (max_pos or 0) + 1
    conn.execute(
        'UPDATE statements SET topic_id = ?, position = ? WHERE id = ?',
        (to_topic_id, next_pos, statement_id),
    )
    conn.commit()
    conn.close()
    return redirect_to_topic_referrer(to_topic_id, statement_id)


@app.route('/move_attachment/<int:attachment_id>', methods=['POST'])
@login_required
def move_attachment(attachment_id):
    """Re-assign an asset to a *different statement* in the SAME topic.

    Semantics (documented in README "Move semantics"):
      - An asset (attachment) lives directly under a statement
        (attachments.statement_id).
      - "Move" here means changing the asset's statement, i.e. moving it
        sideways to a peer statement. Scope is the CURRENT topic only: the
        target statement must belong to the asset's current topic. Cross-topic
        asset moves are out of scope (use the statement Move for that).
    """
    to_statement_id = request.form.get('to_statement_id')
    conn = get_db()
    att = conn.execute('SELECT a.*, t.user_id FROM attachments a JOIN statements s ON a.statement_id = s.id JOIN topics t ON s.topic_id = t.id WHERE a.id = ?', (attachment_id,)).fetchone()
    if not att or att['user_id'] != g.user['id']:
        conn.close()
        flash('Asset not found')
        return redirect(request.referrer or url_for('all_domains'))
    target = conn.execute('SELECT * FROM statements WHERE id = ?', (to_statement_id,)).fetchone()
    if not target:
        conn.close()
        flash('Target statement not found')
        return redirect_to_topic_referrer(att['statement_id'], att['statement_id'])
    if att['statement_id'] == to_statement_id:
        conn.close()
        return redirect_to_topic_referrer(att['statement_id'], att['statement_id'])
    # Scope guard: target statement must live in the asset's current topic.
    current_stmt = conn.execute(
        'SELECT topic_id FROM statements WHERE id = ?', (att['statement_id'],)
    ).fetchone()
    if not current_stmt or current_stmt['topic_id'] != target['topic_id']:
        conn.close()
        flash('An asset can only be moved to another statement in the same topic')
        return redirect_to_topic_referrer(att['statement_id'], att['statement_id'])
    # Append at the end of the destination statement's asset list.
    max_pos = conn.execute(
        'SELECT MAX(position) AS m FROM attachments WHERE statement_id = ?', (to_statement_id,)
    ).fetchone()['m']
    next_pos = (max_pos or 0) + 1
    conn.execute(
        'UPDATE attachments SET statement_id = ?, position = ? WHERE id = ?',
        (to_statement_id, next_pos, attachment_id),
    )
    conn.commit()
    conn.close()
    return redirect_to_topic_referrer(target['topic_id'], to_statement_id)


@app.route('/duplicate_topic/<int:topic_id>', methods=['POST'])
@login_required
def duplicate_topic(topic_id):
    conn = get_db()
    topic = require_topic_owner(conn, topic_id)
    if not topic:
        conn.close()
        flash('Topic not found')
        return redirect(request.referrer or url_for('all_domains'))
    domain_id = topic['domain_id']
    try:
        _copy_topic(conn, topic)
        conn.commit()
    except sqlite3.IntegrityError:
        conn.rollback()
        conn.close()
        flash('Could not duplicate this topic')
        return redirect(url_for('domain', domain_id=domain_id))
    conn.close()
    return redirect(url_for('domain', domain_id=domain_id))


@app.route('/duplicate_statement/<int:statement_id>', methods=['POST'])
@login_required
def duplicate_statement(statement_id):
    conn = get_db()
    stmt = require_statement_owner(conn, statement_id)
    if not stmt:
        conn.close()
        flash('Statement not found')
        return redirect(request.referrer or url_for('all_domains'))
    topic_id = stmt['topic_id']
    # Keep the duplicate in the same topic, appended at the end of the list.
    max_pos = conn.execute(
        'SELECT MAX(position) AS m FROM statements WHERE topic_id = ?', (topic_id,)
    ).fetchone()['m']
    next_pos = (max_pos or 0) + 1
    try:
        cursor = conn.execute(
            "INSERT INTO statements (topic_id, text, position) VALUES (?, 'Copy of ' || ?, ?)",
            (topic_id, stmt['text'], next_pos),
        )
        new_statement_id = cursor.lastrowid
        src_attachments = conn.execute(
            'SELECT * FROM attachments WHERE statement_id = ? ORDER BY position ASC, created_at DESC',
            (statement_id,),
        ).fetchall()
        for src_att in src_attachments:
            _copy_attachment(conn, src_att, new_statement_id)
        conn.commit()
    except sqlite3.IntegrityError:
        conn.rollback()
        conn.close()
        flash('Could not duplicate this statement')
        return redirect_to_topic_referrer(topic_id, statement_id)
    conn.close()
    return redirect_to_topic_referrer(topic_id, new_statement_id)


@app.route('/duplicate_attachment/<int:attachment_id>', methods=['POST'])
@login_required
def duplicate_attachment(attachment_id):
    conn = get_db()
    att = conn.execute('SELECT a.*, t.user_id FROM attachments a JOIN statements s ON a.statement_id = s.id JOIN topics t ON s.topic_id = t.id WHERE a.id = ?', (attachment_id,)).fetchone()
    if not att or att['user_id'] != g.user['id']:
        conn.close()
        flash('Asset not found')
        return redirect(request.referrer or url_for('all_domains'))
    statement_id = att['statement_id']
    topic = conn.execute(
        'SELECT topic_id FROM statements WHERE id = ?', (statement_id,)
    ).fetchone()
    topic_id = topic['topic_id'] if topic else None
    # Keep the duplicate under the same statement, appended at the end.
    max_pos = conn.execute(
        'SELECT MAX(position) AS m FROM attachments WHERE statement_id = ?',
        (statement_id,),
    ).fetchone()['m']
    next_pos = (max_pos or 0) + 1
    try:
        # _copy_attachment reuses the source position; override it afterwards so
        # the copy lands at the end of the same statement's asset list.
        new_att_id = _copy_attachment(conn, att, statement_id)
        conn.execute(
            'UPDATE attachments SET position = ? WHERE id = ?',
            (next_pos, new_att_id),
        )
        conn.commit()
    except sqlite3.IntegrityError:
        conn.rollback()
        conn.close()
        flash('Could not duplicate this asset')
        return redirect_to_topic_referrer(topic_id, statement_id)
    conn.close()
    return redirect_to_topic_referrer(topic_id, statement_id)


@app.route('/duplicate_folder/<int:folder_id>', methods=['POST'])
@login_required
def duplicate_folder(folder_id):
    conn = get_db()
    folder = conn.execute('SELECT f.* FROM folders f WHERE f.id = ?', (folder_id,)).fetchone()
    if not folder or folder['user_id'] != g.user['id']:
        conn.close()
        flash('Folder not found')
        return redirect(request.referrer or url_for('all_domains'))
    domain_id = folder['domain_id']

    def copy_subtree(old_fid, new_parent_id):
        """Recursively copy a folder and its whole subtree, parent before child."""
        src = conn.execute('SELECT * FROM folders WHERE id = ?', (old_fid,)).fetchone()
        cursor = conn.execute(
            "INSERT INTO folders (domain_id, parent_id, name, description, user_id, created_at) "
            "VALUES (?, ?, 'Copy of ' || ?, ?, ?, ?)",
            (domain_id, new_parent_id, src['name'], src['description'], src['user_id'], src['created_at']),
        )
        new_fid = cursor.lastrowid
        src_topics = conn.execute(
            'SELECT * FROM topics WHERE folder_id = ?', (old_fid,)
        ).fetchall()
        for src_topic in src_topics:
            _copy_topic(conn, src_topic, new_fid)
        children = conn.execute(
            'SELECT id FROM folders WHERE parent_id = ?', (old_fid,)
        ).fetchall()
        for child in children:
            copy_subtree(child['id'], new_fid)

    try:
        copy_subtree(folder_id, folder['parent_id'])
        conn.commit()
    except sqlite3.IntegrityError:
        conn.rollback()
        conn.close()
        flash('Could not duplicate this folder')
        return redirect(url_for('domain', domain_id=domain_id))
    conn.close()
    return redirect(url_for('domain', domain_id=domain_id))


@app.route('/update_topic', methods=['POST'])
@login_required
def update_topic():
    topic_id = request.form.get('topic_id')
    name = request.form.get('name', '').strip()
    description = request.form.get('description', '').strip()
    is_public = 1 if request.form.get('is_public') else 0
    if not topic_id or not name:
        flash('Topic name is required')
        return redirect(request.referrer or url_for('all_domains'))
    conn = get_db()
    topic = conn.execute('SELECT t.* FROM topics t WHERE t.id = ?', (topic_id,)).fetchone()
    if not topic or topic['user_id'] != g.user['id']:
        conn.close()
        flash('Topic not found')
        return redirect(request.referrer or url_for('all_domains'))
    conn.execute('UPDATE topics SET name = ?, description = ?, is_public = ? WHERE id = ?', (name, description, is_public, topic_id))
    conn.commit()
    conn.close()
    return redirect(request.referrer or url_for('all_domains'))

@app.route('/delete_topic/<int:topic_id>', methods=['POST'])
@login_required
def delete_topic(topic_id):
    conn = get_db()
    topic = require_topic_owner(conn, topic_id)
    if not topic:
        conn.close()
        flash('Topic not found')
        return redirect(request.referrer or url_for('all_domains'))
    # Collect files before the cascade removes the rows that reference them.
    rows = conn.execute('''
        SELECT a.filename FROM attachments a
        JOIN statements s ON a.statement_id = s.id
        WHERE s.topic_id = ? AND a.filename IS NOT NULL
    ''', (topic_id,)).fetchall()
    conn.execute('DELETE FROM topics WHERE id = ?', (topic_id,))
    conn.commit()
    conn.close()
    for row in rows:
        _remove_upload(row['filename'])
    return redirect(request.referrer or url_for('all_domains'))

@app.route('/uploads/<path:filename>')
def uploaded_file(filename):
    response = send_from_directory(app.config['UPLOAD_FOLDER'], filename)
    # Uploads are user-controlled: never let the browser sniff a different
    # content type, and keep them out of the page's script origin.
    response.headers['X-Content-Type-Options'] = 'nosniff'
    response.headers['Content-Security-Policy'] = "default-src 'none'; sandbox"
    return response


# ---------------------------------------------------------------------------
# Authentication
# ---------------------------------------------------------------------------

@app.route('/signup')
def signup():
    # Pre-fill with server-generated credentials the user may keep or overwrite.
    return render_template(
        'signup.html',
        suggested_username=generate_username(),
        suggested_password=generate_password(),
        next=request.args.get('next', ''),
    )


@app.route('/signup/regenerate')
def signup_regenerate():
    # JSON endpoint so the "regenerate" button can fetch fresh suggestions
    # without a full page reload (keeps creds server-generated).
    return jsonify({'username': generate_username(), 'password': generate_password()})


@app.route('/login')
def login():
    if g.user is not None:
        return redirect(url_for('dashboard'))
    return render_template('login.html', next=request.args.get('next', ''))


@app.route('/auth', methods=['POST'])
def auth():
    """Single create-or-login action used by both the signup and login forms.

    If the username does not yet exist it is created (hashed password) and the
    new account is logged in. If it exists the password is checked and, on
    match, the account is logged in. One button ("Create & Login" / "Login")
    drives both flows; the form submits natively so the browser password
    manager can offer to save the credentials.
    """
    username = (request.form.get('username') or '').strip()
    password = request.form.get('password') or ''
    next_page = request.form.get('next') or request.args.get('next') or ''
    if not username or not password:
        flash('Username and password are required')
        return redirect(url_for('login', next=next_page))

    conn = get_db()
    user = conn.execute('SELECT * FROM users WHERE username = ?', (username,)).fetchone()
    if user is None:
        # New account: create it.
        conn.execute(
            'INSERT INTO users (username, password_hash) VALUES (?, ?)',
            (username, generate_password_hash(password)),
        )
        conn.commit()
        user = conn.execute('SELECT * FROM users WHERE username = ?', (username,)).fetchone()
    else:
        if not check_password_hash(user['password_hash'], password):
            conn.close()
            flash('Incorrect password')
            return redirect(url_for('login', next=next_page))

    conn.close()
    session.clear()
    session['user_id'] = user['id']
    if next_page:
        return redirect(url_for(next_page))
    return redirect(url_for('dashboard'))


@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('index'))


# ---------------------------------------------------------------------------
# Public discovery
# ---------------------------------------------------------------------------

@app.route('/public')
def public_directory():
    """Global directory of every user's public topics and folders."""
    conn = get_db()
    # Public topics with their owning domain and owner username.
    topics = conn.execute('''
        SELECT t.id, t.name, t.description, d.id AS domain_id, d.name AS domain_name,
               u.username AS owner_username
        FROM topics t
        JOIN domains d ON t.domain_id = d.id
        JOIN users u ON t.user_id = u.id
        WHERE t.is_public = 1
        ORDER BY t.created_at DESC
    ''').fetchall()
    # Public folders (containers); their public child topics are shown via /topic.
    folders = conn.execute('''
        SELECT f.id, f.name, f.description, d.id AS domain_id, d.name AS domain_name,
               u.username AS owner_username
        FROM folders f
        JOIN domains d ON f.domain_id = d.id
        JOIN users u ON f.user_id = u.id
        WHERE f.is_public = 1
        ORDER BY f.name
    ''').fetchall()
    conn.close()
    return render_template('public.html', topics=topics, folders=folders)

if __name__ == '__main__':
    init_db()
    debug = os.environ.get('FLASK_DEBUG', '1') == '1'
    app.run(debug=debug, port=10000)
else:
    # Under a WSGI server (gunicorn app:app) the __main__ guard never runs, so
    # the schema migration and bootstrapping must happen at import time or new
    # asset types raise an uncaught IntegrityError on their first insert.
    init_db()
